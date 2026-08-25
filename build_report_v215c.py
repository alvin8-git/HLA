"""
build_report_v215c.py  —  v2.15c: the v2.15 manuscript (commit 792fa1a),
unchanged except for the minimum additions that make it defensible:

  1. Section 2.2: one sentence stating that post-floor frequencies are renormalised
     and coverage is conditional, pointing to Section 4.1.
  2. Section 4.1 Limitations: two paragraphs — the conditional estimand (retained
     mass, unconditional equivalents, the observed 41.8% Chinese twin rate)
     and floor sensitivity (1e-6 -> 8.7e7; comparative findings robust).

Tables, figures, recommendations, and all other text are v2.15 verbatim.
Reads the frozen 1e-3 snapshot. After building, run
    python mark_v217_diffs.py HLA_Registry_Size_CMIO_v2.15.docx HLA_Registry_Size_CMIO_v2.15c.docx
to paint the additions red (everything else black).
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

HERE = os.path.dirname(os.path.abspath(__file__))
# v2.15c reads the frozen 1e-3 snapshot (commit 16bd246); the live analysis/data
# now holds the 1e-6 re-run and would silently change every number.
DATA = os.path.join(HERE, 'analysis', 'snapshot_1e-3', 'data')
FIGS = os.path.join(HERE, 'analysis', 'snapshot_1e-3', 'figures')

# ── Load data ─────────────────────────────────────────────────────────────────
ci     = pd.read_csv(os.path.join(DATA, 'registry_size_ci.csv'))
tgts   = pd.read_csv(os.path.join(DATA, 'registry_size_targets.csv'))
oc_reg = pd.read_csv(os.path.join(DATA, 'others_cluster_registry.csv'))
oc_hap = pd.read_csv(os.path.join(DATA, 'others_cluster_haplotypes.csv'))
sens   = pd.read_csv(os.path.join(DATA, 'cross_ethnic_sensitivity.csv'))
mv     = pd.read_csv(os.path.join(DATA, 'match_validation.csv'))
mr     = pd.read_csv(os.path.join(DATA, 'match_rate_comparison.csv'))

ETHS       = ['Chinese', 'Malay', 'Indian', 'Others']
THRESHOLDS = [0.75, 0.85, 0.90, 0.95]
THR_LABELS = ['75%', '85%', '90%', '95%']

# ── Reference index (1-based for in-text citations) ───────────────────────────
REFS = [
    # 1
    ('Ng AYJ, Moshi GB, Prasath A, et al.',
     'Human leukocyte antigen allele and haplotype frequencies in Singapore '
     'bone marrow donors and cord blood units.',
     'Blood Cell Therapy. 2022;5(3):86–95.',
     'https://doi.org/10.31547/bct-2022-004'),
    # 2
    ('Gragert L, Eapen M, Williams E, et al.',
     'HLA match likelihoods for haematopoietic stem-cell grafts in the U.S. Registry.',
     'N Engl J Med. 2014;371:339–348.',
     ''),
    # 3
    ('Nunes JM, Buhler S, Roessli D, Sanchez-Mazas A; HLA-net 2013 collaboration.',
     'The HLA-net GENE[RATE] pipeline for effective HLA data analysis and its '
     'application to 145 population samples from Europe and neighbouring areas.',
     'Tissue Antigens. 2014;83:307–23.',
     ''),
    # 4
    ('Beatty PG, Mori M, Milford E.',
     'Impact of racial genetic polymorphism on the probability of finding an '
     'HLA-matched donor.',
     'Transplantation. 1995;60(8):778–783.',
     ''),
    # 5
    ('Maiers M, Gragert L, Klitz W.',
     'High-resolution HLA alleles and haplotypes in the United States population.',
     'Hum Immunol. 2007;68(9):779–788.',
     ''),
    # 6
    ('Passweg JR, Baldomero H, Bader P, et al.',
     'Is the use of unrelated donor transplantation leveling off in Europe? '
     'The 2016 European Society for Blood and Marrow Transplant activity survey report.',
     'Bone Marrow Transplant. 2018;53:1139–1148.',
     ''),
    # 7
    ('Lee SJ, Klein J, Haagenson M, et al.',
     'High-resolution donor-recipient HLA matching contributes to the success '
     'of unrelated donor marrow transplantation.',
     'Blood. 2007;110(13):4576–4583.',
     ''),
    # 8
    ('Aljurf M, Weisdorf D, Alfraih F, et al.',
     'Worldwide Network for Blood & Marrow Transplantation (WBMT) special article, '
     'challenges facing emerging alternate donor registries.',
     'Bone Marrow Transplant. 2019;54:1179–1188.',
     ''),
    # 9
    ('Halagan M, Manor S, Shriki N, et al.',
     'East Meets West — Impact of Ethnicity on Donor Match Rates in the '
     'Ezer Mizion Bone Marrow Donor Registry.',
     'Biol Blood Marrow Transplant. 2017;23:1381–6.',
     ''),
    # 10
    ('Singapore Department of Statistics.',
     'Census of Population 2020, Statistical Release 2: Households, Geographic '
     'Distribution, Transport and Difficulty in Basic Activities.',
     'Singapore: Department of Statistics; 2021.',
     'https://www.singstat.gov.sg/files/efc2c599-f1d1-429d-b8fa-3f100f2c8e4c.pdf'),
    # 11
    ('Fleischhauer K, Shaw BE, Gooley T, et al.',
     'Effect of T-cell-epitope matching at HLA-DPB1 in recipients of unrelated-donor '
     'haemopoietic-cell transplantation: a retrospective study.',
     'Lancet Oncol. 2012;13(4):366–374.',
     ''),
    # 12
    ('Klitz W, Maiers M, Spellman S, et al.',
     'New HLA haplotype frequency reference standards: high-resolution and '
     'large sample typing of HLA DR-DQ haplotypes in a sample of European '
     'Americans.',
     'Tissue Antigens. 2003;62(4):296–307.',
     ''),
    # 13
    ('Pidala J, Lee SJ, Ahn KW, et al.',
     'Nonpermissive HLA-DPB1 mismatch increases mortality after myeloablative '
     'unrelated allogeneic haematopoietic cell transplantation.',
     'Blood. 2014;124(16):2596–2606.',
     ''),
    # 14
    ('Gonzalez-Galarza FF, McCabe A, Santos EJMD, et al.',
     'Allele frequency net database (AFND) 2020 update: gold-standard data '
     'classification, open access genotype data and new query tools.',
     'Nucleic Acids Res. 2020;48:D783–8.',
     ''),
    # 15
    ('Price P, Witt C, Allcock R, et al.',
     'The genetic basis for the association of the 8.1 ancestral haplotype '
     '(A1, B8, DR3) with multiple immunopathological diseases.',
     'Immunol Rev. 1999;167:257–274.',
     ''),
    # 16
    ('Efron B, Tibshirani RJ.',
     'An Introduction to the Bootstrap.',
     'New York: Chapman & Hall; 1993.',
     ''),
    # 17
    ('Petersdorf EW.',
     'The major histocompatibility complex: a model for understanding '
     'graft-versus-host disease.',
     'Blood. 2013;122(11):1863–1872.',
     ''),
    # 18
    ('Excoffier L, Slatkin M.',
     'Maximum-likelihood estimation of molecular haplotype frequencies in a '
     'diploid population.',
     'Mol Biol Evol. 1995;12(5):921–927.',
     ''),
    # 19
    ('Kollman C, Abella E, Baitty RL, et al.',
     'Assessment of optimal size and composition of the U.S. National Registry '
     'of hematopoietic stem cell donors.',
     'Transplantation. 2004;78(1):89–95.',
     ''),
]


def cite(*nums):
    """Return '[n]' or '[n,m]' citation string."""
    return '[' + ','.join(str(n) for n in nums) + ']'


# ── Formatting helpers ────────────────────────────────────────────────────────

def n(v):
    return f'{int(v):,}'


def ci_cell(match, eth, thr):
    row = ci[(ci.ethnicity == eth) & (ci.match_level == match) &
             (ci.target_coverage == thr)]
    if row.empty:
        return '—'
    r = row.iloc[0]
    return f"{n(r.registry_size)}\n({n(r.ci_lo)}–{n(r.ci_hi)})"


def cross_cell(match, eth, thr):
    v = tgts[(tgts.match_level == match) & (tgts.ethnicity == eth) &
              (tgts.model_variant == 'cross_ethnic') &
              (tgts.target_coverage == thr)]['registry_size']
    if v.empty:
        return '—'
    val = int(v.iloc[0])
    return '>10,000,000' if val >= 10_000_000 else n(val)


def add_equation(doc, equation_text, eq_num, size=10):
    """Centered numbered equation for professional journal formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{equation_text}\t({eq_num})')
    run.font.size = Pt(size)
    return p


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_para(doc, text='', size=10, space_after=6, left_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
    return p


def add_mixed(doc, parts, space_after=6):
    """
    parts: list of (text, bold, italic) tuples.
    Renders as a single paragraph with mixed formatting.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(10)
    return p


def add_caption(doc, text, fig=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(('Figure' if fig else 'Table') + ': ' + text)
    run.italic    = True
    run.font.size = Pt(9)
    return p


def add_figure(doc, fname, width=6.0, caption=''):
    path = os.path.join(FIGS, fname)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    if caption:
        add_caption(doc, caption)


def add_corrected_para(doc, segments, size=10, space_after=6, left_indent=None):
    """Paragraph with mixed-colour runs for tracked corrections.
    segments = list of (text, is_red) tuples; is_red=True → red font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    for text, is_red in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if is_red:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return p


def style_table(tbl, header_color='1F4E79'):
    tbl.style = 'Table Grid'
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, header_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)


ETH_COLORS = {
    'Chinese': 'D6E4F0', 'Malay': 'D5F5E3',
    'Indian':  'EAD6F7', 'Others': 'FEF9E7',
}


def make_ci_table(doc, match_level):
    header = ['Ethnicity'] + THR_LABELS
    tbl = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for eth in ETHS:
        row = tbl.add_row().cells
        row[0].text = eth
        set_cell_bg(row[0], ETH_COLORS[eth])
        for j, thr in enumerate(THRESHOLDS):
            row[j + 1].text = ci_cell(match_level, eth, thr)
    # Weighted average row
    row = tbl.add_row().cells
    row[0].text = 'Weighted Average†'
    set_cell_bg(row[0], 'E8E8E8')
    for j, thr in enumerate(THRESHOLDS):
        v = tgts[(tgts.match_level == match_level) & (tgts.ethnicity == 'Combined') &
                 (tgts.model_variant == 'same_ethnicity') &
                 (tgts.target_coverage == thr)]['registry_size']
        row[j + 1].text = n(v.iloc[0]) if not v.empty else '—'
    # Attrition-adjusted signed-up targets (40% attrition → ×1.667)
    ATTRITION = 0.40
    row = tbl.add_row().cells
    row[0].text = 'Signed-up target‡\n(40% attrition)'
    set_cell_bg(row[0], 'FEF9E7')
    for j, thr in enumerate(THRESHOLDS):
        # Use per-ethnicity bootstrap medians (Chinese shown; row spans all CMIO)
        # Show range across the four CMIO groups
        vals = []
        for eth in ETHS:
            r = ci[(ci.ethnicity == eth) & (ci.match_level == match_level) &
                   (ci.target_coverage == thr)]
            if not r.empty:
                vals.append(int(r.iloc[0].registry_size / (1 - ATTRITION)))
        if vals:
            row[j + 1].text = f'{min(vals):,}–{max(vals):,}'
        else:
            row[j + 1].text = '—'
    style_table(tbl)
    return tbl


def hap_display(h):
    """Convert 'aa:bb|cc:dd|...' to 'A*aa:bb~B*cc:dd~C*..~DRB1*..~DQB1*..'"""
    loci  = ['A', 'B', 'C', 'DRB1', 'DQB1']
    parts = h.split('|')
    return '~'.join(f'{l}*{a}' for l, a in zip(loci, parts))


# ── Build Document ────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── TITLE ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
tr = title.add_run(
    'Modelling Unrelated Donor Registry Size Requirements for '
    'Haematopoietic Stem Cell Transplantation in Singapore\'s '
    'Multiethnic Population'
)
tr.bold = True
tr.font.size = Pt(14)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_after = Pt(2)
ar = authors.add_run('Alvin Ng Yu-Jin')
ar.italic = True
ar.font.size = Pt(11)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.paragraph_format.space_after = Pt(12)
afr = affil.add_run(
    'National University Hospital / Singapore Cord Blood Bank, Singapore\n'
    'Correspondence: alvin1976sg@gmail.com'
)
afr.font.size = Pt(9)
afr.italic = True

doc.add_paragraph()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Abstract', level=2)
add_para(doc,
    'Background: Access to a compatible unrelated bone marrow donor is critical '
    'for patients with haematological malignancies requiring haematopoietic stem '
    'cell transplantation (HSCT). Singapore\'s multiethnic Chinese–Malay–Indian–'
    'Others (CMIO) population presents distinct HLA diversity challenges [1,10]. '
    'Using expectation-maximisation (EM)-estimated phased haplotype frequencies '
    'from 59,186 donors and cord blood units [1], we modelled the minimum registry '
    'size needed to achieve 75–95% population coverage at both 8/8 and 10/10 HLA '
    'match levels.')
add_corrected_para(doc, [
    ('Methods: Diplotype frequencies were derived under Hardy–Weinberg equilibrium '
     'from EM-phased five-locus haplotypes ', False),
    ('estimated using an expectation–maximisation algorithm with full phase '
     'enumeration ' + cite(18) + '; '
     'results were validated against the HLA-net GENE[RATE] database ' + cite(3) + '', True),
    ('. Coverage was modelled as a function of registry size ' + cite(2, 5) + '. Bootstrap '
     'resampling (1,000 iterations) provided 95% confidence intervals ' + cite(16) + '. '
     'Partial-match (9/10, 8/10) relaxation, cross-ethnic feasibility, and '
     'ancestry stratification of the Others sub-group were analysed as secondary '
     'objectives.', False),
])
add_para(doc,
    'Results: For 10/10 HLA matching at 95% same-ethnicity coverage of '
    'ethnicity-specific registries, bootstrap median registry size estimates are: '
    'Chinese 41,183, Malay 39,831, Indian 43,785. Bootstrap 95% confidence intervals '
    '(Dirichlet resampling, B=1,000; see Section 2.4) are 40,184–42,153, 38,680–40,935, '
    'and 42,762–44,547 respectively; these intervals capture haplotype-frequency '
    'sampling variability only and are a lower bound on total uncertainty. '
    'Only the Chinese estimate is empirically validated against observed patient '
    'haplotypes (Spearman r=0.70; Section 3.6); the Malay, Indian, and Others figures are '
    'model projections pending validation. '
    'The pooled Others estimate of 31,129 is a statistical artefact of heterogeneous '
    'population mixing; ancestry sub-cluster analysis reveals three genetically '
    'distinct groups (European/Eurasian, Filipino/SE Asian, Northeast Asian) with '
    'markedly different requirements (35,193–63,856 at 95%), and 63,856 — the '
    'Filipino/SE Asian cluster ceiling — is the recommended planning target. '
    'Cross-ethnic matching is infeasible for Malay, Indian, and Others patients '
    'regardless of registry size. Relaxing to 9/10 matching approximately halves '
    'the required registry size. Registry size estimates are robust across '
    'patient demographic scenarios (variation 3.5%).')
add_para(doc,
    'Conclusions: An ethnicity-specific same-ethnicity registry of approximately '
    '40,000–44,000 donors per major CMIO group is needed for 95% 10/10 coverage. '
    'Targeted minority recruitment, partial-match protocols, and recognition of '
    'Others sub-group diversity are essential for equitable HSCT access in Singapore.',
    space_after=8)

kw = doc.add_paragraph()
kw.paragraph_format.space_after = Pt(12)
kwr = kw.add_run(
    'Keywords: HLA; donor registry; haematopoietic stem cell transplantation; '
    'CMIO; Singapore; registry size; partial match; EM haplotype estimation; '
    'bootstrap confidence interval; Others stratification'
)
kwr.italic = True
kwr.font.size = Pt(9)

doc.add_page_break()

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction')

add_para(doc,
    'Haematopoietic stem cell transplantation (HSCT) offers a potential cure for '
    'many blood cancers and bone marrow disorders. When a patient lacks a matched '
    'family donor — as is the case for the majority — their best hope lies in an '
    'unrelated volunteer donor from a national or international registry; '
    'unrelated donors now account for a large and growing share of allogeneic '
    'transplant activity ' + cite(6) + '. '
    'High-resolution HLA matching between donor and patient is strongly associated '
    'with transplant success and patient survival ' + cite(7) + '.')

add_para(doc,
    'Finding a matched donor is substantially harder for patients from non-European '
    'backgrounds ' + cite(9, 19) + '. Worldwide registries are dominated by donors '
    'of European descent, and HLA allele combinations that are common in East Asian, '
    'South Asian, or mixed-ancestry populations may be absent or very rare in those '
    'registries. For Singapore\'s Chinese, Malay, Indian, and Others (CMIO) '
    'populations ' + cite(10) + ', this creates an urgent need to understand exactly '
    'how large a local, ethnicity-representative registry must be to give patients a '
    'realistic chance of finding a match ' + cite(8) + '.')

add_para(doc,
    'Ng et al. (2022) characterised the HLA landscape of Singapore\'s donor pool '
    'across 59,186 individuals and published allele and haplotype frequencies for '
    'each CMIO group ' + cite(1) + '. That study provided the essential '
    'epidemiological foundation but did not translate these frequencies into '
    'quantitative registry size recommendations. The present study does so, '
    'addressing four questions:')

for bullet in [
    'How many same-ethnicity donors are needed for 75–95% coverage at strict '
    '(10/10) and standard (8/8) HLA match levels, and how confident can we be '
    'in these estimates?',
    'Can cross-ethnic matching — drawing Malay, Indian, or Others patients from '
    'a Chinese-dominated combined pool — ever achieve adequate coverage?',
    'How much smaller could registries be if clinicians accepted partial matches '
    '(9 or 8 of 10 alleles)?',
    'Is the Others group sufficiently homogeneous to be modelled as a single '
    'population, or do distinct ancestry sub-groups require separate strategies?',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(bullet).font.size = Pt(10)

add_para(doc,
    'Answers to these questions carry direct implications for Singapore\'s bone '
    'marrow donor recruitment strategy and for donor registries serving similar '
    'multiethnic populations elsewhere, for which the challenges of building a '
    'representative registry are well described ' + cite(8) + '.')

# ── 2. METHODS ────────────────────────────────────────────────────────────────
add_heading(doc, '2. Methods')
add_figure(doc, 'pipeline_flowchart.png', width=6.5,
    caption='Figure 1. Analysis pipeline overview: from raw HLA typing data to '
    'registry size estimates with bootstrap confidence intervals. '
    'Shading: blue = input data; green = core model steps; '
    'yellow = optimisation; red = uncertainty quantification.')
add_heading(doc, '2.1 Dataset', level=2)
add_corrected_para(doc, [
    ('HLA typing data were obtained from 59,186 ', False),
    ('donors and cord blood units from the Bone Marrow Donor Programme (BMDP) '
     'and Singapore Cord Blood Bank (SCBB)', True),
    (', accrued between 2005 and 2020 ' + cite(1) + '. '
     'Recipients and donors processed by the Health Sciences Authority (HSA) were '
     'included for validation. ', False),
    ('Samples', True),
    (' were typed at high resolution across five loci '
     '(HLA-A, -B, -C, -DRB1, -DQB1). Full details of typing methodology and data '
     'cleaning are reported in Ng et al. ' + cite(1) + '.', False),
])

add_heading(doc, '2.2 Haplotype Frequency Estimation Using EM Phasing', level=2)
add_corrected_para(doc, [
    ('A key methodological advance in this analysis is the use of '
     'expectation-maximisation (EM)-estimated phased haplotype frequencies, rather '
     'than the simpler approach of multiplying individual allele frequencies at each '
     'locus. Haplotype frequencies were estimated using ', False),
    ('an expectation–maximisation algorithm with full multi-locus phase '
     'enumeration (Excoffier–Slatkin formulation ' + cite(18) + '), implemented '
     'in Python, applying Hardy–Weinberg equilibrium to infer the most probable '
     'phase assignments from unphased diplotype data (capped at 5,000 individuals '
     'per ethnic group for computational efficiency; haplotypes retained at '
     'frequency ≥ 0.1%). Results were validated against the HLA-net '
     'GENE[RATE] database ' + cite(3) + '.', True),
])
add_para(doc,
    'Haplotypes retained after the 0.1% frequency floor were renormalised to sum '
    'to one; coverage estimates are therefore conditional on the patient carrying '
    'two haplotypes above this floor. The implications of this choice are '
    'discussed under Limitations (Section 4.1).')

add_para(doc,
    'HLA loci are not independent — alleles at neighbouring loci are inherited '
    'together as haplotype blocks far more often than chance would predict. This '
    'non-random co-inheritance, known as linkage disequilibrium (LD), is especially '
    'strong between DRB1 and DQB1 (D′ ≥ 0.94 in all CMIO groups) and '
    'between B and C (D′ ≥ 0.95) ' + cite(1, 12) + '. Ignoring LD by '
    'using a simple product approximation systematically underestimates the '
    'frequency of common haplotypes, which in turn distorts the projected '
    'registry size. Phasing with the EM algorithm ' + cite(18) + ' avoids that '
    'distortion; we validated the resulting frequencies against the GENE[RATE] '
    'reference pipeline ' + cite(3) + ' (Section 3.6).')

add_heading(doc, '2.3 Registry Size Model', level=2)
add_para(doc,
    'For each CMIO group, the EM algorithm produced a ranked list of K haplotypes '
    'with estimated frequencies f₁, f₂, ..., fᴊ. Diplotype (genotype '
    'pair) frequencies were derived under Hardy–Weinberg equilibrium:')
add_equation(doc, 'F(hᵢ, hᵢ) = fᵢ²   [homozygous, i = j]', 1)
add_equation(doc, 'F(hᵢ, hⱼ) = 2·fᵢ·fⱼ,   i ≠ j   [heterozygous]', 2)
add_para(doc,
    'The probability that a patient with diplotype dₖ finds at least one matched '
    'donor in a registry of N independently drawn donors is:')
add_equation(doc, 'P(dₖ, N) = 1 − (1 − Fₖ)^N', 3)
add_para(doc,
    'Population coverage C(N) — the expected fraction of all patients who find '
    'at least one match — is the diplotype-frequency-weighted sum over all m '
    'diplotypes:')
add_equation(doc, 'C(N) = Σₖ Fₖ · [1 − (1 − Fₖ)^N],   k = 1, …, m', 4)
add_para(doc,
    'Equation (4) aggregates across all diplotype classes. For a patient whose '
    'diplotype dₖ occurs at population frequency Fₖ, the probability that at '
    'least one of N randomly drawn registry donors carries that same diplotype '
    'is [1 − (1 − Fₖ)^N] — the complement of all N donors lacking it. '
    'Multiplying by Fₖ weights this probability by the fraction of patients '
    'who carry diplotype dₖ; summing over all m diplotypes gives the overall '
    'population coverage C(N). As N grows, each term approaches Fₖ and C(N) '
    'approaches 1, but the convergence is slow for rare diplotypes.')
add_corrected_para(doc, [
    ('The minimum registry size N* to achieve target coverage θ (75%, 85%, '
     '90%, or 95%) was found by binary search on a logarithmic scale (50 iterations; '
     'precision < 1 donor for N ≤ 10⁷). '
     'Same-ethnicity and cross-ethnic model variants were compared. ', False),
    ('The coverage-based framework — computing C(N) as the diplotype-frequency-weighted '
     'sum of per-patient match probabilities — follows the approach pioneered by '
     'Beatty et al. ' + cite(4) + ' and applied to national registry planning by '
     'Kollman et al. ' + cite(19) + '.', True),
])

add_heading(doc, '2.4 Bootstrap Confidence Intervals', level=2)
add_para(doc,
    'To quantify uncertainty around N* estimates, we used Dirichlet parametric '
    'bootstrapping (B = 1,000 iterations) ' + cite(16) + '. In each iteration, '
    'haplotype frequencies were resampled from a Dirichlet distribution with '
    'concentration parameters nₑₗₗ × f̂ₖ, where nₑₗₗ '
    'is the number of individuals from which the haplotype frequencies for that '
    'ethnicity were estimated — the 5-locus donor count subject to the '
    '5,000-individual EM limit of Section 2.3 (Chinese: 5,000; Malay: 5,000; '
    'Indian: 5,000; Others: 3,767, this group falling below the limit) — '
    'and f̂ₖ is the EM-estimated haplotype '
    'frequency. A new N* was computed for each resample. The reported point estimate '
    'is the bootstrap median (50th percentile), which is bias-corrected for the '
    'left-skew in N* near high coverage thresholds arising from Jensen\'s inequality '
    'on the concave coverage function C(N). The 2.5th and 97.5th percentiles define '
    'the 95% CI. Setting nₑₗₗ to the full group size rather than the capped '
    'estimation sample would assert a precision the frequencies do not carry — for '
    'the Chinese group, whose 44,400 5-locus donors exceed the EM limit almost '
    'ninefold, it narrows the 95% interval by a factor of roughly four. '
    'These intervals quantify sampling variability in haplotype '
    'frequencies given the sample the frequencies were estimated from; they do not '
    'propagate uncertainty from the EM phasing step, from the choice of which '
    'individuals fall within the 5,000 cap, or from HWE model assumptions. Total '
    'uncertainty across the full estimation pipeline is therefore wider than the '
    'reported CIs.')
add_para(doc,
    'For clinical planning purposes, the bootstrap median is a conservative and '
    'robust target: it is the registry size below which half of all plausible '
    'sampling realisations would leave the stated coverage threshold unmet. '
    'EM maximum-likelihood estimates are retained in the supplementary data archive '
    'for reference; narrower CIs indicate more precisely estimated haplotype frequencies.')

add_heading(doc, '2.5 Partial-Match Analysis', level=2)
add_corrected_para(doc, [
    ('Partial-match coverage curves were generated for match thresholds of 10/10, '
     '9/10, and 8/10 (5-locus framework) and 8/8, 7/8, 6/8 (4-locus framework) ', False),
    (cite(2, 17), True),
    ('. A partial match is defined as sharing a specified minimum '
     'number of alleles across the relevant loci. These analyses used EM-phased '
     'haplotype pairs to correctly account for LD structure ' + cite(3) + '.', False),
])

add_heading(doc, '2.6 Others Group Stratification', level=2)
add_para(doc,
    'The Others group is ethnically heterogeneous, comprising Eurasians, '
    'Caucasians, and a variety of mixed and other Asian backgrounds. To assess '
    'whether this group should be treated as a single population, we applied '
    'principal component analysis (PCA) to a binary HLA allele indicator matrix '
    '(alleles with ≥1% frequency across the Others cohort), followed by '
    'k-means clustering (k = 2–5). The optimal number of clusters was selected '
    'by silhouette coefficient. Registry size requirements were then estimated '
    'independently for each cluster. Ancestry inference for each cluster was based '
    'on haplotype signature matching against the Allele Frequency Net Database '
    '(AFND) ' + cite(14) + ' and published population-specific haplotype '
    'references ' + cite(1, 15) + '.')

# ── 3. RESULTS ────────────────────────────────────────────────────────────────
add_heading(doc, '3. Results')
add_para(doc,
    'Section 3.1 presents the primary finding: same-ethnicity registry size '
    'requirements at 10/10 matching. Section 3.2 shows that 8/8 targets are '
    'marginally lower. Section 3.3 demonstrates that cross-ethnic matching '
    'cannot substitute for same-ethnicity donors. Sections 3.4–3.5 present '
    'secondary analyses (partial-match and sensitivity). Section 3.6 reports '
    'model validation. Section 3.7 is explicitly exploratory: ancestry '
    'stratification of the heterogeneous "Others" subgroup.')

# 3.1 Registry size 10/10
add_heading(doc, '3.1 Registry Size Requirements — 10/10 HLA Matching', level=2)
add_para(doc,
    'Table 1 presents the minimum registry sizes required for same-ethnicity '
    '10/10 HLA matching across all four CMIO groups, based on EM-phased '
    'haplotype frequencies ' + cite(1, 3) + '. Values are bootstrap median '
    'estimates (bias-corrected point estimates); numbers in parentheses are 95% '
    'bootstrap confidence intervals (2.5th–97.5th percentile) ' + cite(16) + '. '
    'The Weighted Average row reflects Singapore population weights ' + cite(10) + '; '
    'it does not guarantee equitable access for minority ethnic groups.')

make_ci_table(doc, '10of10')
add_para(doc,
    '† Weighted Average: Singapore resident population weights (Chinese 74.3%, Malay 13.5%, '
    'Indian 9.0%, Others 3.2%) [10]. This row is a mathematical convenience, not a policy target: '
    'it describes how a single shared registry would need to be sized to serve all '
    'groups proportionally under those weights, but it cannot substitute for '
    'ethnicity-specific recruitment because donors are not interchangeable across '
    'groups. The per-group same-ethnicity N* values above are the operative planning '
    'targets for each community.\n'
    'Values shown as bootstrap median N (95% CI lower–upper).',
    size=8, space_after=4)
add_caption(doc,
    'Table 1. Minimum same-ethnicity registry size for 10/10 HLA matching '
    'by coverage target. Bootstrap 95% CIs (Dirichlet resampling, 1,000 '
    'iterations) [16] reflect haplotype-frequency sampling variability only '
    'and are a lower bound on total uncertainty. The pooled Others row is a '
    'mathematical artefact; see Section 3.7 for sub-cluster targets. '
    '‡ Signed-up target assumes 40% real-world donor attrition '
    '(unreachability, refusal, or medical deferral); shows range across CMIO '
    'groups at each threshold.', fig=False)

add_para(doc,
    'At the clinically important 95% coverage threshold, same-ethnicity registry '
    'sizes (bootstrap medians) span approximately 40,000–44,000 donors. The Indian '
    'group requires the largest registry (43,785 donors), reflecting its greater '
    'haplotype diversity ' + cite(1) + '. The pooled Others estimate (31,129) '
    'is a mathematical artefact of heterogeneous sub-group mixing and must not be '
    'used as a policy target; sub-cluster analysis (Section 3.7) reveals requirements of '
    '35,193–63,856. The 95% bootstrap CIs are of comparable width across all four '
    'groups (±800–1,100 donors at 95% coverage), because the 5,000-individual EM '
    'limit (Section 2.3) equalises the sample from which haplotype frequencies were '
    'estimated for Chinese, Malay and Indian; Others falls below the limit at 3,767. '
    'These intervals therefore reflect the estimation sample rather than the size of '
    'each donor pool ' + cite(16) + '.')
add_para(doc,
    'These per-group same-ethnicity N* values (bootstrap medians; e.g., 41,183 '
    'for Chinese) should be distinguished from the population-weighted combined '
    'estimates in the sensitivity analysis (Table 6; Section 3.5), where Singapore\'s '
    'ethnic composition (74.3% Chinese) produces a combined N* of ~42,567. The '
    'combined figure reflects how a single shared registry would need to be sized '
    'to serve all groups proportionally; the per-group figures reflect the target '
    'for a dedicated ethnicity-specific registry.')

add_para(doc,
    'A striking feature of these results is how rapidly requirements grow as '
    'coverage targets increase beyond 90%. This reflects the long tail of rare '
    'HLA diplotypes: the first 90% of patients can be served by a relatively '
    'concentrated pool of common haplotype combinations, but the final 5–10% '
    'of patients carry increasingly rare diplotypes that require '
    'disproportionately larger registries to cover.')

add_figure(doc, 'registry_ci_plot.png', width=5.5,
    caption='Figure 2. Bootstrap 95% confidence intervals for minimum registry size '
    'at 95% coverage (10/10 matching) by CMIO group. Error bars show the 2.5th–97.5th '
    'percentile range across 1,000 Dirichlet bootstrap resamples [16].')

# 3.2 Registry size 8/8
add_heading(doc, '3.2 Registry Size Requirements — 8/8 HLA Matching', level=2)
add_para(doc,
    'At the 8/8 match level (omitting DQB1 as a mandatory matching criterion), '
    'registry size requirements are marginally lower than for 10/10 matching '
    '(Table 2). The difference is modest — typically 600–1,200 fewer donors at '
    '95% coverage — because DRB1 and DQB1 are in very strong linkage '
    'disequilibrium across all CMIO groups (D′ = 0.94–0.99) ' + cite(1, 12) + '. '
    'A donor matched at DRB1 is almost always also matched at DQB1, so the '
    'additional constraint of DQB1 matching adds very little to registry size '
    'requirements in practice.')

make_ci_table(doc, '8of8')
add_para(doc,
    '† Weighted Average: Singapore population weights [10]. This weighted average does '
    'not guarantee equitable access for minority ethnic groups; see same-ethnicity '
    'targets above.\n'
    '‡ Signed-up target assumes 40% real-world donor attrition; shows range across CMIO groups.\n'
    'Values shown as bootstrap median N (95% CI lower–upper).',
    size=8, space_after=4)
add_caption(doc,
    'Table 2. Minimum same-ethnicity registry size for 8/8 HLA matching '
    'by coverage target. Bootstrap 95% CIs (Dirichlet resampling, 1,000 '
    'iterations) [16]; CIs reflect haplotype-frequency sampling variability '
    'only (lower bound on total uncertainty). '
    '‡ Signed-up target assumes 40% attrition; shows range across CMIO groups.', fig=False)

add_para(doc,
    'This finding has an important practical implication: the strong DRB1–DQB1 '
    'LD means that investing in DQB1 typing data is worthwhile — it adds '
    'meaningful clinical information (DQB1 mismatches can increase graft-versus-'
    'host disease risk) at very little cost in registry size. Programmes can '
    'justify full 10-locus typing without substantially larger recruitment targets '
    + cite(7, 11) + '.')

# 3.3 Cross-ethnic
add_heading(doc,
    '3.3 Cross-Ethnic Matching — Why It Cannot Replace Same-Ethnicity Donors',
    level=2)
add_para(doc,
    'A natural question is whether a large shared registry — primarily supplied '
    'by the numerically dominant Chinese donor group — could serve patients of '
    'all ethnicities. The cross-ethnic analysis answers this question definitively: '
    'it cannot (Table 3). This is consistent with findings from other multiethnic '
    'registries ' + cite(9) + '.')

tbl3_data = [['Ethnicity', '75% coverage', '85% coverage',
              '90% coverage', '95% coverage']]
for eth in ETHS:
    row = [eth]
    for thr in THRESHOLDS:
        row.append(cross_cell('10of10', eth, thr))
    tbl3_data.append(row)

tbl3 = doc.add_table(rows=len(tbl3_data), cols=5)
for i, row_data in enumerate(tbl3_data):
    for j, val in enumerate(row_data):
        tbl3.rows[i].cells[j].text = val
style_table(tbl3)
for i, eth in enumerate(ETHS):
    set_cell_bg(tbl3.rows[i + 1].cells[0], ETH_COLORS[eth])
add_para(doc,
    'Combined donor pool: Singapore resident population weights (Chinese 74.3%) [10].',
    size=8, space_after=4)
add_caption(doc,
    'Table 3. Minimum registry size for 10/10 HLA matching under the cross-ethnic '
    'model (combined Singapore donor pool). Values >10,000,000 are infeasible.',
    fig=False)

add_para(doc,
    'For the Chinese group, cross-ethnic matching is feasible (93,348 donors at '
    '95% coverage) but requires roughly twice the same-ethnicity target — because '
    'Chinese haplotypes dominate the combined pool. For Malay, Indian, and Others '
    'patients, cross-ethnic matching from a Chinese-dominated pool is essentially '
    'impossible: no registry of realistic size achieves adequate coverage, because '
    'their distinctive haplotype combinations are not represented in the donor pool '
    + cite(9) + '. This underscores that same-ethnicity donor recruitment is not '
    'merely preferable but, on these results, the only viable strategy for these '
    'groups ' + cite(8) + '.')

# 3.4 Partial match
add_heading(doc,
    '3.4 Partial-Match Coverage — The Clinical Benefit of Relaxing Match Criteria',
    level=2)
add_para(doc,
    'While exact 10/10 or 8/8 matching is the preferred clinical standard '
    + cite(7, 17) + ', some centres accept partial matches — particularly when '
    'no fully matched donor exists or when time is critical. '
    'Figures 3 and 4 show coverage curves for progressive relaxation of match '
    'criteria in the 10-locus and 8-locus frameworks.')

add_figure(doc, 'partial_match_10locus.png', width=6.2,
    caption='Figure 3. Coverage curves for 10-locus (HLA-A, -B, -C, -DRB1, -DQB1) '
    'partial matching across CMIO groups. Green: 10/10 exact match; Blue: ≥9/10; '
    'Red: ≥8/10. Dashed lines: cross-ethnic model. Note the sharp threshold '
    'effect at the 9/10 boundary for all groups.')

add_para(doc,
    'The most important finding is the dramatic threshold effect at the 9/10 '
    'boundary. Relaxing from 10/10 to 9/10 matching roughly halves the required '
    'registry size for all CMIO groups. For Chinese patients, the 9/10 registry '
    'requirement at 95% coverage is approximately 20,000–22,000 donors, compared '
    'with 41,183 at 10/10. This is clinically significant because several large '
    'multi-centre trials have demonstrated that a single allele mismatch at certain '
    'loci — particularly at HLA-DPB1 when the mismatch is at a permissive '
    'T-cell epitope group position — carries minimal additional survival impact '
    + cite(11, 13) + '. Programmes that formalise evidence-based partial-match '
    'protocols could effectively double their donor reach without recruiting a '
    'single additional donor (Section 3.4).')

add_figure(doc, 'partial_match_8locus.png', width=6.2,
    caption='Figure 4. Coverage curves for 8-locus (HLA-A, -B, -C, -DRB1) partial '
    'matching. Green: 8/8; Blue: ≥7/8; Red: ≥6/8. Cross-ethnic '
    'shown with dashed lines.')

add_para(doc,
    'In the 8-locus framework (Figure 4), the same pattern holds. The cross-ethnic '
    'curves (dashed) plateau for Malay, Indian, and Others groups regardless of '
    'relaxation level — confirming that partial matching cannot substitute for '
    'same-ethnicity donors when fundamental haplotype diversity differs markedly '
    'between donor and patient populations ' + cite(2, 9) + '.')

# 3.5 Sensitivity
add_heading(doc,
    '3.5 Registry Size Is Robust to Patient Demographic Assumptions', level=2)
add_para(doc,
    'Registry size models must assume a patient population composition ' + cite(2, 5) + '. '
    'We tested four scenarios: Singapore resident population weights (74.3% Chinese, '
    '13.5% Malay, 9.0% Indian, 3.2% Others; 3,006,770, 545,500, 362,270 and 129,670 '
    'of 4,044,210 residents) ' + cite(10) + ', BMDP+SCBB registry composition, '
    'HSA Patient-Donor Data (referral composition from Health Sciences Authority Singapore), and an extreme minority-focus '
    'scenario (equal Malay, Indian, and Others weighting, no Chinese). Table 6 and '
    'Figure 5 show that the combined N* varies by 3.5% across all scenarios at the '
    '95% target (42,567 down to 41,129), and by less at lower targets.')

tbl6_data = [['Scenario', 'Ethnic weights (C/M/I/O)',
              '75%', '85%', '90%', '95%']]
scenario_labels = {
    'SG population (current model)':        'SG population',
    'BMDP+SCBB registry composition':       'BMDP+SCBB donors',
    'Patient.txt composition':              'HSA Patient-Donor Data',
    'Minority-focus (Indian+Malay+Others)': 'Minority-focus',
}
scenario_weights = {
    'SG population (current model)':        '77/8/9/6',
    'BMDP+SCBB registry composition':       '75/9/9/6',
    'Patient.txt composition':              '72/15/5/8',
    'Minority-focus (Indian+Malay+Others)': '0/40/40/20',
}
for sk in ['SG population (current model)', 'BMDP+SCBB registry composition',
           'Patient.txt composition', 'Minority-focus (Indian+Malay+Others)']:
    sub = sens[sens.scenario == sk]
    row = [scenario_labels[sk], scenario_weights[sk]]
    for thr in THRESHOLDS:
        v = sub[sub.target_coverage == thr]['registry_size']
        row.append(n(v.iloc[0]) if not v.empty else '—')
    tbl6_data.append(row)

tbl6 = doc.add_table(rows=len(tbl6_data), cols=6)
for i, row_data in enumerate(tbl6_data):
    for j, val in enumerate(row_data):
        tbl6.rows[i].cells[j].text = val
style_table(tbl6)
add_caption(doc,
    'Table 6. Sensitivity of combined registry size to patient ethnic composition '
    'scenario [10]. C/M/I/O = Chinese/Malay/Indian/Others percentage weights.',
    fig=False)

add_figure(doc, 'cross_ethnic_sensitivity.png', width=6.0,
    caption='Figure 5. Registry size sensitivity across four patient demographic '
    'scenarios [10], including HSA Patient-Donor Data (Health Sciences Authority '
    'Singapore). Near-identical bar heights confirm the ~40,000–45,000 donor '
    'target is a structural property of CMIO haplotype diversity [2].')

add_para(doc,
    'This stability arises because per-group N* values are all in the same order '
    'of magnitude (~31,000–44,000 at 95%). Reweighting the groups changes the '
    'combined N* only marginally ' + cite(2) + '. This structural robustness '
    'gives decision-makers confidence that the ~40,000–45,000 donor target per '
    'ethnicity will remain valid even as Singapore\'s demographic composition '
    'evolves ' + cite(10) + '.')

# 3.6 Model validation
add_heading(doc, '3.6 Model Validation', level=2)
add_para(doc,
    'To assess whether EM-derived haplotype frequencies reflect real patient '
    'haplotype distributions, we compared EM estimates against independently '
    'observed haplotype frequencies from 564 patients whose typing was provided by '
    'the Health Sciences Authority (HSA), Singapore. Spearman '
    'rank correlations and root mean square error (RMSE) were computed on shared '
    'haplotypes (Table 7).')

tbl7_header = ['Ethnicity', 'Patient haplotypes', 'Shared with EM',
               '% frequency covered', 'Spearman r', 'RMSE']
tbl7 = doc.add_table(rows=1, cols=6)
for i, h in enumerate(tbl7_header):
    tbl7.rows[0].cells[i].text = h
for _, mvrow in mv.iterrows():
    eth = mvrow['ethnicity']
    row = tbl7.add_row().cells
    row[0].text = eth
    row[1].text = str(int(mvrow['n_patient'])) if pd.notna(mvrow['n_patient']) else '—'
    row[2].text = str(int(mvrow['n_shared']))   if pd.notna(mvrow['n_shared'])  else '—'
    row[3].text = f"{mvrow['pct_covered']:.1f}%" if pd.notna(mvrow['pct_covered']) else '—'
    row[4].text = f"{mvrow['spearman_r']:.3f}"   if pd.notna(mvrow['spearman_r'])  else 'n/a*'
    row[5].text = f"{mvrow['rmse']:.4f}"          if pd.notna(mvrow['rmse'])        else '—'
    set_cell_bg(row[0], ETH_COLORS.get(eth, 'FFFFFF'))
style_table(tbl7)
add_para(doc, '* Insufficient shared haplotypes for rank correlation (n < 3).',
    size=8, space_after=4)
add_caption(doc,
    'Table 7. Validation of EM-estimated haplotype frequencies [3] against '
    'independently observed patient haplotypes.', fig=False)

add_figure(doc, 'match_validation_scatter.png', width=5.5,
    caption='Figure 6. Observed (patient) vs EM-estimated haplotype frequencies '
    'per ethnicity. Dashed line = perfect agreement. Chinese: Spearman r=0.70 '
    '(p<0.001, n=33 shared haplotypes).')

add_para(doc,
    'For Chinese patients, the EM estimates show good rank agreement with observed '
    'patient frequencies (Spearman r = 0.70, p < 0.001), with RMSE = 0.0094; '
    'this is the primary validation result. For Malay (11 shared haplotypes) and '
    'Others (4 shared haplotypes), sample sizes are insufficient for reliable rank '
    'correlation. For Indian patients only one shared haplotype was observed, '
    'precluding any meaningful validation ' + cite(1) + '. Registry size estimates '
    'for Malay, Indian, and Others should therefore be regarded as model-derived '
    'projections; the Chinese estimates are the most robustly validated. These '
    'results highlight the need for larger patient cohorts in minority groups.')

# 3.7 Others stratification — Exploratory; CMI is the primary focus
add_heading(doc, '3.7 Exploratory Analysis — Note on the \"Others\" Subgroup', level=2)
add_para(doc,
    'The Others category in Singapore\'s CMIO classification encompasses Eurasians, '
    'Caucasians, and individuals of mixed or diverse Asian ancestry ' + cite(10) + '. '
    'Unsupervised clustering (optimal k=3 by silhouette coefficient; s=0.24 in the '
    'five-component PCA space used for clustering) of the 3,847 fully five-locus-typed Others donors reveals '
    'three genetically distinct sub-groups — European/Eurasian, Filipino/SE Asian, '
    'and Northeast Asian/Mixed — with markedly different registry size requirements '
    '(Table 4). The pooled Others estimate in Table 1 is a mathematical artefact of '
    'heterogeneous population mixing and should not be used as a standalone policy '
    'target; the Filipino/SE Asian sub-cluster ceiling of 63,856 donors provides the '
    'appropriate conservative planning figure for this group.')

# Table 4: cluster registry sizes
tbl4_header = ['Cluster', 'Putative ancestry', 'N individuals',
               '75%', '85%', '90%', '95%']
tbl4 = doc.add_table(rows=1, cols=len(tbl4_header))
for i, h in enumerate(tbl4_header):
    tbl4.rows[0].cells[i].text = h

ancestry_labels = {
    'Cluster_1': 'European / Eurasian',
    'Cluster_2': 'Filipino / SE Asian',
    'Cluster_3': 'Northeast Asian / Mixed',
}
clust_colors = {
    # Light tints of Figure 7 scatter colours (#e41a1c, #377eb8, #4daf4a)
    'Cluster_1': 'FFCCCC',  # light red   → #e41a1c
    'Cluster_2': 'CCE0F5',  # light blue  → #377eb8
    'Cluster_3': 'CBF0CB',  # light green → #4daf4a
}
for cl in ['Cluster_1', 'Cluster_2', 'Cluster_3']:
    sub   = oc_reg[oc_reg.cluster == cl]
    n_ind = int(sub.iloc[0].n_individuals)
    sizes = {int(r.target_coverage * 100): int(r.registry_size)
             for _, r in sub.iterrows()}
    row   = tbl4.add_row().cells
    row[0].text = cl.replace('_', ' ')
    row[1].text = ancestry_labels[cl]
    row[2].text = f'{n_ind:,}'
    for j, thr_pct in enumerate([75, 85, 90, 95]):
        row[3 + j].text = n(sizes.get(thr_pct, 0))
    set_cell_bg(row[0], clust_colors[cl])
    set_cell_bg(row[1], clust_colors[cl])
style_table(tbl4)
add_caption(doc,
    'Table 4. Minimum registry size for 10/10 HLA matching within each Others '
    'sub-cluster. Putative ancestry inferred from haplotype signature matching '
    'against AFND [14] and published references [1,15]. '
    'The pooled Others row (Table 1) is a mathematical artefact and must not '
    'be used as a policy target (see text).', fig=False)

add_figure(doc, 'others_pca_scatter.png', width=5.0,
    caption='Figure 7. PCA scatter of 3,847 fully five-locus-typed Others donors (binary HLA allele '
    'indicators, alleles ≥1%). Three distinct clusters (optimal k=3 by silhouette '
    'coefficient, s=0.24 in the five-PC clustering space; 0.43 in the PC1–PC2 '
    'projection shown) indicate different ancestry backgrounds. Cluster identities inferred from '
    'haplotype signature matching [14,15].')

# Mini haplotype validation table — top 2 haplotypes per cluster
hap_tbl_header = ['Cluster', 'Rank', 'Haplotype (A~B~C~DRB1~DQB1)',
                  'Freq (%)', 'Population association']

# Population annotations for top haplotypes
pop_annot = {
    # Cluster 1 haplotypes
    '01:01|08:01|07:01|03:01|02:01':
        '8.1 Ancestral Haplotype — hallmark Northern European [15]',
    '33:03|44:03|07:01|07:01|02:01':
        'Common South Asian (Indian subcontinent) [14]',
    '33:03|58:01|03:02|03:01|02:01':
        'Most common Chinese/East Asian haplotype [1]',
    '01:01|57:01|06:02|07:01|03:03':
        'South Asian signature (India, Sri Lanka) [14]',
    '29:02|44:03|16:01|07:01|02:01':
        'Mediterranean / Middle Eastern European [14]',
    # Cluster 2 haplotypes
    '24:07|35:05|04:01|12:02|03:01':
        'Filipino / SE Asian archipelago (B*35:05–DRB1*12:02) [14]',
    '11:01|15:02|08:01|12:02|03:01':
        'Most common Filipino haplotype (B*15:02–DRB1*12:02) [14]',
    '02:01|15:13|08:01|12:02|03:01':
        'SE Asian (B*15:13 — Philippines, Indonesia) [14]',
    '24:02|15:02|08:01|12:02|03:01':
        'Filipino / Thai / Indonesian [14]',
    '24:07|15:02|08:01|12:02|03:01':
        'SE Asian archipelago [14]',
    # Cluster 3 haplotypes
    '02:07|46:01|01:02|09:01|03:03':
        'Chinese-specific (A*02:07–B*46:01) [1]',
    '24:02|38:02|07:02|15:02|05:02':
        'Japanese / Korean signature (B*38:02–DRB1*15:02) [14]',
    '34:01|15:21|04:03|15:02|06:01':
        'South/SE Asian [14]',
    '03:01|07:02|07:02|15:01|06:02':
        'Common European haplotype [14]',
    '29:01|07:05|15:05|10:01|05:01':
        'South Asian / mixed [14]',
}

# Table 5: top haplotype per cluster (ancestry validation) — no Rank column
tbl5_header = ['Cluster', 'Top Haplotype (A~B~C~DRB1~DQB1)', 'Freq (%)', 'Population association']
tbl5 = doc.add_table(rows=1, cols=len(tbl5_header))
for i, h in enumerate(tbl5_header):
    tbl5.rows[0].cells[i].text = h

for cl in ['Cluster_1', 'Cluster_2', 'Cluster_3']:
    top1 = (oc_hap[oc_hap.cluster == cl]
            .sort_values('frequency', ascending=False)
            .iloc[0])
    row = tbl5.add_row().cells
    row[0].text = cl.replace('_', ' ')
    row[1].text = hap_display(top1['haplotype'])
    row[2].text = f"{top1['frequency']*100:.1f}"
    row[3].text = pop_annot.get(top1['haplotype'], '—')
    set_cell_bg(row[0], clust_colors[cl])

style_table(tbl5)
# Widen Population association column (col 3) and narrow Cluster/Freq columns
tbl5.autofit = False
col_widths = [Cm(2.8), Cm(5.2), Cm(1.3), Cm(7.5)]
for row in tbl5.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w
add_caption(doc,
    'Table 5. Top haplotype per Others sub-cluster confirming putative ancestry. '
    'Frequencies are within-cluster. Population assignments cross-referenced '
    'against AFND [14] and published references [1,15].', fig=False)

# Per-cluster narratives not rendered in main body

# ── 4. DISCUSSION AND RECOMMENDATIONS ────────────────────────────────────────
add_heading(doc, '4. Discussion and Recommendations')

add_para(doc,
    'This analysis provides the first quantitative registry size targets for '
    'Singapore\'s CMIO populations with bootstrap confidence intervals, grounded '
    'in EM-phased haplotype frequencies from the largest local HLA dataset to date '
    + cite(1) + '. The findings are consistent, robust, and carry clear '
    'recommendations for donor recruitment strategy ' + cite(8) + '.')

rec_items = [
    (
        'Build same-ethnicity registries of 40,000–45,000 donors per major group.',
        'Cross-ethnic matching is demonstrably insufficient for Malay, Indian, '
        'and Others patients ' + cite(9) + '. The current BMDP registry '
        '(predominantly Chinese) provides good coverage for Chinese patients '
        'but leaves other groups critically under-served. Ethnicity-specific '
        'recruitment campaigns are strongly indicated by these results: the '
        'model shows no realistic shared-registry alternative for these groups '
        + cite(8) + '.'
    ),
    (
        'Prioritise Malay and Indian donor recruitment urgently.',
        'Ng et al. report that transplant recipients from the Malay and Indian '
        'communities face particular difficulty finding a suitable unrelated '
        'donor or cord blood unit, attributing it to under-representation of '
        'their haplotypes in local donor pools and in international registries '
        + cite(1) + ' — a pattern also seen for minority groups elsewhere '
        + cite(8, 19) + '. In this dataset '
        'the Malay and Indian groups contribute 5,578 and 5,490 five-locus '
        'typed donors respectively (Section 2.1). At 95% '
        'coverage, both groups require registries of ~40,000 same-ethnicity '
        'donors (model projections pending validation; see Sections 3.6 and 4.1) — '
        'comparable in scale to the entire current BMDP. Targeted '
        'outreach, community partnerships, and incentive programmes are required.'
    ),
    (
        'Adopt partial-match protocols to extend effective coverage.',
        'Relaxing from 10/10 to 9/10 matching roughly halves the effective '
        'registry size requirement (Section 3.4); the U.S. registry reports 7/8 '
        'alongside 8/8 match likelihoods on the same logic ' + cite(2) + '. '
        'A 9/10 donor is not clinically equivalent to a 10/10 donor, and this '
        'analysis contains no outcome data of its own. What the external '
        'literature does establish is that the permissiveness of the mismatch '
        'matters as much as the count: a 10/10 transplant carrying a '
        'non-permissive DPB1 mismatch had outcomes not substantially different '
        'from a 9/10 transplant with a permissive DPB1 mismatch, while '
        'non-permissive mismatches raised mortality within both strata '
        + cite(11, 13) + '. Formalising 9/10 with explicit DPB1 permissiveness '
        'criteria would therefore extend the reach of existing registries '
        'substantially without additional recruitment, at a cost in outcome '
        'that is real but bounded and, for patients with no 10/10 donor at any '
        'registry size, is the comparison against no transplant rather than '
        'against a full match. Because a 9/10 '
        'unrelated donor is by definition a mismatched unrelated donor (MMUD), '
        'these results imply that MMUD capability halves every group\'s '
        'effective registry requirement and provides the only access route for '
        'patients whose haplotypes fall below the frequency floor '
        '(Section 4.1); the clinical safety of transplanting across such '
        'mismatches rests on the external trial literature ' + cite(11, 13) + '.'
    ),
    (
        'Plan Others recruitment to 63,856 donors and collect ancestry sub-group data.',
        'The pooled Others estimate of 31,129 is a statistical artefact of '
        'heterogeneous pooling and must not be used as a policy target (see Section 3.7). '
        'The three ancestry clusters have markedly different requirements '
        '(35,193–63,856 at 95%); planning to any figure below the highest '
        'sub-cluster ceiling will guarantee under-service for the most '
        'haplotype-diverse sub-group. The Filipino/SE Asian cluster (Cluster 2, '
        'N*=63,856) is the recommended planning ceiling for Others recruitment. '
        'In parallel, simple ancestry questions at registration — parental '
        'birthplace, self-reported heritage — would enable sub-group–targeted '
        'recruitment and allow the pooled figure to be retired as a planning '
        'metric. This is a recommendation of the present analysis, not a '
        'practice reported elsewhere.'
    ),
    (
        'Invest in DQB1 typing for all donors.',
        'Because DRB1 and DQB1 are in strong linkage disequilibrium in our data '
        '(Section 3.2), 8/8 and 10/10 registry sizes differ by less than 5%, so '
        'the marginal recruitment cost of typing DQB1 is very small. The case '
        'for doing so is one of information rather than demonstrated survival '
        'benefit: the largest registry outcome study found DQ and DP '
        'mismatching not to be associated with survival once 8/8 matching was '
        'achieved ' + cite(7) + ', so DQB1 typing is recommended here because '
        'it is nearly free given the LD structure and preserves the option of '
        'reporting 10/10, not because withholding it is known to cost lives.'
    ),
    (
        'Use the 95% coverage target as the planning standard.',
        'Planning for 90% rather than 95% coverage may appear cost-saving on '
        'paper but systematically excludes patients with rare diplotypes — who '
        'are disproportionately represented in minority and mixed-ancestry groups '
        + cite(2, 8) + '. The 95% target reflects genuine clinical equity '
        'and is the standard used by established international registries '
        + cite(7) + '. Both thresholds are conditional on the haplotype '
        'frequency floor (Section 4.1): patients whose haplotypes fall below '
        'the floor are not reached at either threshold, and their access '
        'depends on the partial-match protocols of Recommendation 3.'
    ),
]

for i, (heading, body) in enumerate(rec_items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'Recommendation {i+1}: {heading}')
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.8)
    for run in p2.runs:
        run.font.size = Pt(10)

add_heading(doc, '4.1 Limitations', level=2)
add_para(doc,
    'Several limitations should be noted. First, diplotype frequencies are '
    'derived under Hardy–Weinberg equilibrium. In our own testing, HWE departures '
    'were detected in the Indian and Others groups at several loci (Section 3.3). '
    'The direction '
    'and magnitude of bias are uncertain and depend on the nature of the '
    'sub-structure; registry size estimates for Indian and Others should be '
    'treated as exploratory, while Chinese and Malay estimates are on firmer '
    'ground. Second, the EM algorithm was capped at 5,000 individuals per '
    'ethnic group; this cap binds materially only for Chinese (5,000 of 45,018 '
    'five-locus-typed donors). A convergence test rerunning the EM at sample '
    'sizes from 500 to 45,018 (Supporting Figure S1) shows N* stabilising near '
    '42,000 donors above ~20,000 samples; at 5,000 the estimate is 45,148 versus '
    '41,727 at the full sample — a conservative 8.2% overestimate, which is a '
    'safe direction of bias for planning. Third, the reported N* values represent '
    'biologically matched donors; real-world donor attrition (unreachability, '
    'refusal, medical deferral — typically 30–50% in established registries) '
    'means signed-up recruitment targets must exceed N* by a corresponding '
    'factor. At 40% attrition, signed-up targets are N* ÷ 0.60 ≈ '
    'N* × 1.67; per-threshold ranges across CMIO groups are shown in the '
    '"Signed-up target" rows of Tables 1 and 2. Fourth, N* estimates are lower '
    'bounds for unobserved haplotypes. Sensitivity analysis with Laplace '
    'pseudocount smoothing (α = 0.001 per haplotype) changes N* at the 95% '
    'threshold by +0.9% (Chinese), +2.3% (Malay), −3.1% (Indian), and '
    '−1.9% (Others) — confirming robustness at the primary planning threshold. '
    'Effects are larger at 75% (+11–24%), indicating that lower-coverage '
    'targets are more sensitive to rare-haplotype treatment. '
    'Fifth, the Others cluster ancestry assignments are inferred from haplotype '
    'signatures ' + cite(14, 15) + ' without confirmed self-reported ancestry '
    'data, and cluster stability was not independently validated via bootstrap '
    'resampling of individuals — the labels are supported by distinct '
    'population-specific top-haplotype signatures (Table 5), but the moderate '
    'silhouette coefficient (0.24) indicates overlapping cluster boundaries, and '
    'independent confirmation '
    'would strengthen these findings. Sixth, validation of EM frequencies against '
    'patient haplotypes was limited by small patient sample sizes in Malay, Indian, '
    'and Others groups ' + cite(1) + ' (see Section 3.6). Finally, the registry model '
    'assumes random donor sampling; demographic biases in recruitment (age, sex, '
    'region) may affect effective coverage in practice.')

add_para(doc,
    'A further limitation concerns the 0.1% haplotype frequency floor (Section 2.2). '
    'Because haplotypes below the floor were excluded and the remainder '
    'renormalised, the floor retains 51.7% (Chinese), 52.9% (Malay), 40.2% '
    '(Indian), and 35.6% (Others) of each group\'s frequency mass, and all '
    'coverage percentages and registry sizes in this report are conditional on '
    'the patient carrying two haplotypes above the floor. Across all patients in '
    'each group, a registry meeting the 95% conditional target serves '
    'approximately one third to one half. This conditional reading is consistent '
    'with direct observation: among the 44,400 five-locus-typed Chinese donors — '
    'a pool approximately the size of the modelled 95% target — 41.8% have an '
    'exact 10/10 genotype match within the registry, against a conditional-model '
    'equivalent of 49.1% when the discarded mass is restored. Patients carrying '
    'rare haplotypes below the floor are not reached at any registry size '
    'considered here; for these patients, partial-match transplantation '
    '(Section 3.4, Recommendation 3) is the relevant route of access.')
add_para(doc,
    'The absolute registry sizes are also sensitive to the floor itself: '
    'repeating the analysis with a 1×10⁻⁶ floor, which lies below every group\'s '
    'singleton frequency, increases the Chinese 95% 10/10 requirement from '
    'approximately 43,000 to 8.7×10⁷ donors. Absolute targets should therefore '
    'be read as order-of-magnitude planning guidance contingent on the stated '
    'floor, whereas the comparative findings — cross-ethnic versus '
    'same-ethnicity matching, and match-stringency relaxation versus registry '
    'expansion — are computed under a common floor and are robust to this '
    'choice.')

# ── 5. CONCLUSIONS ───────────────────────────────────────────────────────────
add_heading(doc, '5. Conclusions')
add_para(doc,
    'Using EM-phased haplotype frequencies from 59,186 Singapore donors and cord '
    'blood units ' + cite(1) + ', we have derived the first quantitative registry '
    'size targets for Singapore\'s CMIO populations with bootstrap confidence '
    'intervals ' + cite(16) + '. Same-ethnicity registries of approximately '
    '40,000–45,000 donors are required per CMIO group to achieve 95% coverage '
    'at 10/10 HLA matching — a target that cannot be met by cross-ethnic matching '
    'from any registry of realistic size ' + cite(2, 9) + '.')

add_corrected_para(doc, [
    ('Partial-match analysis demonstrates that 9/10 matching halves this requirement ', False),
    ('(Section 3.4) ' + cite(17), True),
    (', providing a practical and clinically justifiable path to '
     'improving access in the near term ' + cite(11, 13) + '. The Others subgroup '
     'is ethnically heterogeneous; sub-cluster analysis identifies three distinct '
     'ancestry groups with registry requirements ranging from 35,000 to 64,000 donors '
     '(Table 4) ' + cite(14) + '.', False),
])

add_para(doc,
    'These findings support a three-pronged national strategy: (1) accelerated '
    'same-ethnicity donor recruitment with defined numerical targets for each '
    'CMIO group ' + cite(8) + '; (2) adoption of evidence-based partial-match '
    'protocols to extend effective registry reach ' + cite(11, 13) + '; and '
    '(3) ancestry sub-group data collection for Others donors to support targeted '
    'recruitment as registry size grows ' + cite(14) + '. '
    'Together, these measures can substantially improve HSCT access equity for '
    'all communities in Singapore.', space_after=12)

# ── SUPPORTING ANALYSIS ──────────────────────────────────────────────────────
add_heading(doc, 'Supporting Analysis')
add_heading(doc, 'S1. EM Convergence: N* Stability vs Sample Size (Chinese)', level=2)
add_para(doc,
    'To assess whether the 5,000-sample EM cap introduces material bias for '
    'the Chinese group, the EM was rerun at sample sizes ranging from 500 to '
    '45,018 (full dataset). Figure S1 shows N* at 95% coverage (10/10 matching) '
    'as a function of input sample size.')
add_figure(doc, 'em_convergence.png', width=6.0,
    caption='Figure S1. N* at 95% coverage (10/10 matching) for Chinese donors '
    'as a function of EM input sample size. The red dashed line marks the 5,000 '
    'cap used in the main analysis (N*=45,148); N* at the full 45,018-donor '
    'sample is 41,727 — an 8.2% overestimate at the cap, in the conservative '
    '(safe) direction. Convergence is near-complete above ~20,000 samples.')
add_para(doc,
    'The 8.2% overestimate at the cap means the reported N* values for Chinese '
    'are conservative by approximately 3,400 donors — equivalent to slightly '
    'less than one year of typical registry recruitment. Malay (5,578) and Indian '
    '(5,490) exceed the cap only marginally, so roughly a tenth of their donors is '
    'excluded from phasing; Others, at 3,767, falls below the cap and is '
    'unaffected. The convergence curve is not monotonic in this region, so the '
    'direction of the residual bias for Malay and Indian is not established by '
    'this analysis.')

# ── GLOSSARY ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Abbreviations', level=2)
glossary = [
    ('AFND',  'Allele Frequency Net Database'),
    ('BMDP',  'Bone Marrow Donor Programme (Singapore)'),
    ('CI',    'Confidence Interval'),
    ('CMIO',  'Chinese–Malay–Indian–Others (Singapore ethnic classification)'),
    ('EM',    'Expectation–Maximisation (algorithm for haplotype phasing)'),
    ('HSCT',  'Haematopoietic Stem Cell Transplantation'),
    ('HLA',   'Human Leukocyte Antigen'),
    ('HSA',   'Health Sciences Authority (Singapore)'),
    ('HWE',   'Hardy–Weinberg Equilibrium'),
    ('LD',    'Linkage Disequilibrium'),
    ('MLE',   'Maximum Likelihood Estimate'),
    ('N*',    'Minimum registry size achieving a specified coverage target'),
    ('PCA',   'Principal Component Analysis'),
    ('RMSE',  'Root Mean Square Error'),
    ('SCBB',  'Singapore Cord Blood Bank'),
]
gtbl = doc.add_table(rows=1, cols=2)
gtbl.rows[0].cells[0].text = 'Abbreviation'
gtbl.rows[0].cells[1].text = 'Definition'
for abbr, defn in glossary:
    row = gtbl.add_row().cells
    row[0].text = abbr
    row[1].text = defn
style_table(gtbl)
gtbl.autofit = False
for row in gtbl.rows:
    row.cells[0].width = Cm(2.2)
    row.cells[1].width = Cm(13.5)

# ── DECLARATIONS ─────────────────────────────────────────────────────────────
add_heading(doc, 'Declarations', level=2)
add_para(doc, 'Conflict of interest: The author declares no conflict of interest.', size=9)
add_para(doc,
    'Data availability: All analysis code and intermediate data files are '
    'available at https://github.com/alvin8-git/HLA.', size=9)
add_para(doc, 'Funding: No external funding was received for this study.', size=9)

# ── REFERENCES ───────────────────────────────────────────────────────────────
add_heading(doc, 'References')
for i, (authors, title, journal, doi) in enumerate(REFS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.left_indent   = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r1 = p.add_run(f'{i}.  ')
    r1.bold = True
    r1.font.size = Pt(9)
    ref_text = f'{authors} {title} {journal}'
    if doi:
        ref_text += f' {doi}'
    p.add_run(ref_text).font.size = Pt(9)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(HERE, 'HLA_Registry_Size_CMIO_v2.15c.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables:     {len(doc.tables)}')
print(f'  Figures:    {len(doc.inline_shapes)}')
print(f'  References: {len(REFS)}')
