#!/usr/bin/env python3
"""Render the BMT-targeted manuscript markdown to .docx (no pandoc available).

Usage: python build_paper_docx.py [input.md]
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else 'HLA_Registry_Size_CMIO_BoneMarrowTransplantation.md'
OUT = SRC.rsplit('.', 1)[0] + '.docx'

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(1)

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')


def rich(p, text):
    """Add text to paragraph honouring **bold**, *italic*, `code`."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Courier New'; r.font.size = Pt(10)
        else:
            p.add_run(part.replace('\\*', '*'))


def table(rows):
    hdr = [c.strip() for c in rows[0].strip('|').split('|')]
    body = [[c.strip() for c in r.strip('|').split('|')] for r in rows[2:]]
    t = doc.add_table(rows=1 + len(body), cols=len(hdr))
    t.style = 'Table Grid'
    for j, h in enumerate(hdr):
        cell = t.cell(0, j).paragraphs[0]
        rich(cell, h)
        for r in cell.runs:
            r.bold = True
            r.font.size = Pt(9)
    for i, row in enumerate(body, start=1):
        for j, v in enumerate(row[:len(hdr)]):
            cell = t.cell(i, j).paragraphs[0]
            rich(cell, v)
            for r in cell.runs:
                r.font.size = Pt(9)
    doc.add_paragraph()


lines = open(SRC).read().split('\n')
i = 0
in_front = False
while i < len(lines):
    ln = lines[i]
    if ln.strip() == '---' and i == 0:
        in_front = True; i += 1; continue
    if in_front:
        if ln.strip() == '---':
            in_front = False
        i += 1; continue
    if ln.strip() == '---':
        i += 1; continue

    if ln.startswith('|') and i + 1 < len(lines) and set(lines[i + 1].replace('|', '').strip()) <= set('-: '):
        blk = []
        while i < len(lines) and lines[i].startswith('|'):
            blk.append(lines[i]); i += 1
        table(blk); continue

    m = re.match(r'^(#{1,4})\s+(.*)', ln)
    if m:
        lvl, txt = len(m.group(1)), m.group(2)
        if lvl == 1:
            p = doc.add_paragraph()
            r = p.add_run(txt); r.bold = True; r.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
        else:
            h = doc.add_heading(level=min(lvl, 3))
            hr = h.add_run(txt)
            hr.font.name = 'Times New Roman'
            hr.font.color.rgb = RGBColor(0, 0, 0)
            hr.font.size = Pt({2: 13, 3: 11.5, 4: 11}[min(lvl, 4)])
        i += 1; continue

    if re.match(r'^\s*[-*]\s+', ln):
        p = doc.add_paragraph(style='List Bullet')
        rich(p, re.sub(r'^\s*[-*]\s+', '', ln)); i += 1; continue

    if re.match(r'^\d+\.\s+', ln):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        rich(p, ln); i += 1; continue

    if ln.strip():
        buf = [ln.strip()]          # always consume at least this line — else infinite loop
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(('#', '|')) \
                and not re.match(r'^\s*[-*]\s+|^\d+\.\s', lines[i]):
            buf.append(lines[i].strip()); i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rich(p, ' '.join(buf))
        continue
    i += 1

doc.save(OUT)
print(f"Saved {OUT}")
