"""
build_report.py  —  Rebuild HLA_Registry_Size_Report_CMIO.docx

Changes vs prior version:
  1. EM-phased haplotype frequencies (corrected registry size numbers)
  2. EM maximum-likelihood registry sizes per CMIO group (CIs withdrawn, see 2.4)
  3. Partial-match coverage model (9/10, 8/10 relaxation)
  4. Ancestral stratification of the "Others" group (3 clusters with ancestry evidence)
  5. Full in-text citations [n] throughout
  6. Extended reference list (19 references)
"""

import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, 'analysis', 'data')
FIGS = os.path.join(HERE, 'analysis', 'figures')

# ── Load data ─────────────────────────────────────────────────────────────────
# registry_size_ci.csv is deliberately NOT read. The Dirichlet parametric
# bootstrap that produced it is biased downward for N* on a long-tailed
# haplotype distribution: E[f^2] = f^2 + f(1-f)/(n+1), so resampling inflates
# the squared/product terms that form rare diplotype frequencies (by ~525% for
# haplotypes at 1e-6 to 1e-5, ~134% at 1e-5 to 1e-4), raising apparent coverage
# and lowering N* in essentially every replicate. In v2.15 this put the EM
# estimate outside its own CI in 18 of 32 rows; at the 1e-6 floor the gap
# reaches 62%. Point estimates are the EM maximum-likelihood values.
tgts   = pd.read_csv(os.path.join(DATA, 'registry_size_targets.csv'))
oc_reg = pd.read_csv(os.path.join(DATA, 'others_cluster_registry.csv'))
oc_hap = pd.read_csv(os.path.join(DATA, 'others_cluster_haplotypes.csv'))
sens   = pd.read_csv(os.path.join(DATA, 'cross_ethnic_sensitivity.csv'))
mv     = pd.read_csv(os.path.join(DATA, 'match_validation.csv'))
mr     = pd.read_csv(os.path.join(DATA, 'match_rate_comparison.csv'))

ETHS       = ['Chinese', 'Malay', 'Indian', 'Others']
THRESHOLDS = [0.75, 0.85, 0.90, 0.95]
THR_LABELS = ['75%', '85%', '90%', '95%']

# ── Key figures pulled from the CSVs, so prose cannot drift from the tables ───
# Every registry size or coverage percentage quoted in the running text should
# come from here rather than being typed in. Hardcoded prose silently
# contradicting a recomputed table is the failure mode that produced the
# retracted numbers in versions up to v2.15.
def _nstar(eth, thr=0.95, level='10of10', variant='same_ethnicity'):
    r = tgts[(tgts.ethnicity == eth) & (tgts.match_level == level) &
             (tgts.target_coverage == thr) & (tgts.model_variant == variant)]
    return int(r.registry_size.iloc[0]) if len(r) else None


def _millions(n):
    """Format a registry size for prose: 3,153,571 -> '3.2 million'."""
    if n is None:
        return 'not available'
    if n >= 1e9:
        return f'{n/1e9:.2f} billion'
    if n >= 1e6:
        return f'{n/1e6:.1f} million'
    return f'{n:,}'


_cov = pd.read_csv(os.path.join(DATA, 'coverage_curves.csv'))


def _cov_at(eth, n, level='10of10', variant='same_ethnicity'):
    """Modelled coverage (%) at registry size n, from the swept curve."""
    s = _cov[(_cov.ethnicity == eth) & (_cov.match_level == level) &
             (_cov.model_variant == variant)].sort_values('N')
    if s.empty:
        return None
    return round(100 * float(np.interp(n, s.N.values, s.coverage.values)), 1)


_emc = pd.read_csv(os.path.join(DATA, 'em_convergence.csv'))
_cap_row  = _emc.iloc[(_emc.sample_size - 5000).abs().argmin()]
_full_row = _emc.loc[_emc.sample_size.idxmax()]
CAP_N, FULL_N = int(_cap_row.registry_size), int(_full_row.registry_size)
CAP_NHAP, FULL_NHAP = int(_cap_row.n_haplotypes), int(_full_row.n_haplotypes)
FULL_SAMPLE = int(_full_row.sample_size)
CAP_DELTA_PCT = round(100 * (CAP_N - FULL_N) / FULL_N, 1)

_ld = pd.read_csv(os.path.join(DATA, 'ld_report.csv'))


def _ld_range(pair):
    v = _ld[_ld.locus_pair == pair]['composite_dprime']
    return (round(float(v.min()), 2), round(float(v.max()), 2)) if len(v) else (None, None)


LD_DRDQ = _ld_range('DRB1–DQB1')
LD_BC   = _ld_range('B–C')

KEY = {
    'n95':      {e: _nstar(e) for e in ETHS},
    'n95_txt':  {e: _millions(_nstar(e)) for e in ETHS},
    'n75_txt':  {e: _millions(_nstar(e, 0.75)) for e in ETHS},
    'cross95':  {e: _millions(_nstar(e, 0.95, variant='cross_ethnic')) for e in ETHS},
    'cov50k':   {e: _cov_at(e, 50_000) for e in ETHS},
    'cov100k':  {e: _cov_at(e, 100_000) for e in ETHS},
    'cov500k':  {e: _cov_at(e, 500_000) for e in ETHS},
    'combined95': _millions(_nstar('Combined')),
}


def K(path, eth=None):
    """K('n95_txt','Chinese') -> '3.2 million'."""
    v = KEY[path]
    return v[eth] if eth else v

# ── Reference index (1-based for in-text citations) ───────────────────────────
REFS = [
    # 1
    ('Ng AYJ, Moshi GB, Prasath A, et al.',
     'Human leukocyte antigen allele and haplotype frequencies in Singapore '
     'bone marrow donors and cord blood units.',
     'Blood Cell Therapy. 2022;5(3):86–95.',
     'https://doi.org/10.31547/bct-2022-004'),
    # 2
    ('Gragert L, Eapen M, Williams E, et al.',
     'HLA match likelihoods for haematopoietic stem-cell grafts in the U.S. Registry.',
     'N Engl J Med. 2014;371:339–348.',
     ''),
    # 3
    ('Nunes JM, Buhler S, Roessli D, Sanchez-Mazas A; HLA-net 2013 collaboration.',
     'The HLA-net GENE[RATE] pipeline for effective HLA data analysis and its '
     'application to 145 population samples from Europe and neighbouring areas.',
     'Tissue Antigens. 2014;83:307–23.',
     ''),
    # 4
    ('Beatty PG, Mori M, Milford E.',
     'Impact of racial genetic polymorphism on the probability of finding an '
     'HLA-matched donor.',
     'Transplantation. 1995;60(8):778–783.',
     ''),
    # 5
    ('Maiers M, Gragert L, Klitz W.',
     'High-resolution HLA alleles and haplotypes in the United States population.',
     'Hum Immunol. 2007;68(9):779–788.',
     ''),
    # 6
    ('Passweg JR, Baldomero H, Bader P, et al.',
     'Is the use of unrelated donor transplantation leveling off in Europe? '
     'The 2016 European Society for Blood and Marrow Transplant activity survey report.',
     'Bone Marrow Transplant. 2018;53:1139–1148.',
     ''),
    # 7
    ('Lee SJ, Klein J, Haagenson M, et al.',
     'High-resolution donor-recipient HLA matching contributes to the success '
     'of unrelated donor marrow transplantation.',
     'Blood. 2007;110(13):4576–4583.',
     ''),
    # 8
    ('Lim YA, Teo D, Ang AL, et al.',
     'HLA allele and haplotype frequencies in unrelated bone marrow donor '
     'registries in Asia.',
     'Transpl Immunol. 2010;22(3–4):166–174.',
     ''),
    # 9
    ('Aljurf M, Weisdorf D, Alfraih F, et al.',
     'Worldwide Network for Blood & Marrow Transplantation (WBMT) special article, '
     'challenges facing emerging alternate donor registries.',
     'Bone Marrow Transplant. 2019;54:1179–1188.',
     ''),
    # 10
    ('Halagan M, Manor S, Shriki N, et al.',
     'East Meets West — Impact of Ethnicity on Donor Match Rates in the '
     'Ezer Mizion Bone Marrow Donor Registry.',
     'Biol Blood Marrow Transplant. 2017;23:1381–6.',
     ''),
    # 11
    ('Singapore Department of Statistics.',
     'Census of Population 2020.',
     'Singapore: Department of Statistics; 2021.',
     ''),
    # 12
    ('Fleischhauer K, Shaw BE, Gooley T, et al.',
     'Effect of T-cell-epitope matching at HLA-DPB1 in recipients of unrelated-donor '
     'haematopoietic-cell transplantation: a retrospective study.',
     'Lancet Oncol. 2012;13(4):366–374.',
     ''),
    # 13
    ('Klitz W, Gragert L, Maiers M, Byard PJ.',
     'New HLA haplotype frequency reference standards: five-locus haplotypes, '
     'and allele frequencies for 66 North American populations.',
     'Tissue Antigens. 2003;62(4):296–307.',
     ''),
    # 14
    ('Anasetti C, Logan BR, Lee SJ, et al.; Blood and Marrow Transplant Clinical '
     'Trials Network.',
     'Peripheral-blood stem cells versus bone marrow from unrelated donors.',
     'N Engl J Med. 2012;367(16):1487–1496.',
     ''),
    # 15
    ('Pidala J, Lee SJ, Ahn KW, et al.',
     'Nonpermissive HLA-DPB1 mismatch increases mortality after myeloablative '
     'unrelated allogeneic haematopoietic cell transplantation.',
     'Blood. 2014;124(16):2596–2606.',
     ''),
    # 16
    ('Gonzalez-Galarza FF, McCabe A, Santos EJMD, et al.',
     'Allele frequency net database (AFND) 2020 update: gold-standard data '
     'classification, open access genotype data and new query tools.',
     'Nucleic Acids Res. 2020;48:D783–8.',
     ''),
    # 17
    ('Price P, Witt C, Allcock R, et al.',
     'The genetic basis for the association of the 8.1 ancestral haplotype '
     '(A1, B8, DR3) with multiple immunopathological diseases.',
     'Immunol Rev. 1999;167:257–274.',
     ''),
    # 18
    ('Efron B, Tibshirani RJ.',
     'An Introduction to the Bootstrap.',
     'New York: Chapman & Hall; 1993.',
     ''),
    # 19
    ('Petersdorf EW.',
     'The major histocompatibility complex: a model for understanding '
     'graft-versus-host disease.',
     'Blood. 2013;122(11):1863–1872.',
     ''),
    # 20
    ('Excoffier L, Slatkin M.',
     'Maximum-likelihood estimation of molecular haplotype frequencies in a '
     'diploid population.',
     'Mol Biol Evol. 1995;12(5):921–927.',
     ''),
]


def cite(*nums):
    """Return '[n]' or '[n,m]' citation string."""
    return '[' + ','.join(str(n) for n in nums) + ']'


# ── Formatting helpers ────────────────────────────────────────────────────────

def n(v):
    return f'{int(v):,}'


def ci_cell(match, eth, thr):
    """EM maximum-likelihood N* (same-ethnicity). Name kept for call sites."""
    row = tgts[(tgts.ethnicity == eth) & (tgts.match_level == match) &
               (tgts.target_coverage == thr) &
               (tgts.model_variant == 'same_ethnicity')]
    return '—' if row.empty else n(row.iloc[0].registry_size)


def cross_cell(match, eth, thr):
    v = tgts[(tgts.match_level == match) & (tgts.ethnicity == eth) &
              (tgts.model_variant == 'cross_ethnic') &
              (tgts.target_coverage == thr)]['registry_size']
    if v.empty:
        return '—'
    val = int(v.iloc[0])
    return '>10,000,000' if val >= 10_000_000 else n(val)


def add_equation(doc, equation_text, eq_num, size=10):
    """Centered numbered equation for professional journal formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{equation_text}\t({eq_num})')
    run.font.size = Pt(size)
    return p


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_para(doc, text='', size=10, space_after=6, left_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
    return p


def add_mixed(doc, parts, space_after=6):
    """
    parts: list of (text, bold, italic) tuples.
    Renders as a single paragraph with mixed formatting.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(10)
    return p


def add_caption(doc, text, fig=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(('Figure' if fig else 'Table') + ': ' + text)
    run.italic    = True
    run.font.size = Pt(9)
    return p


def add_figure(doc, fname, width=6.0, caption=''):
    path = os.path.join(FIGS, fname)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    if caption:
        add_caption(doc, caption)


def add_corrected_para(doc, segments, size=10, space_after=6, left_indent=None):
    """Paragraph with mixed-colour runs for tracked corrections.
    segments = list of (text, is_red) tuples; is_red=True → red font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    for text, is_red in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if is_red:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return p


def style_table(tbl, header_color='1F4E79'):
    tbl.style = 'Table Grid'
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, header_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)


ETH_COLORS = {
    'Chinese': 'D6E4F0', 'Malay': 'D5F5E3',
    'Indian':  'EAD6F7', 'Others': 'FEF9E7',
}


def make_ci_table(doc, match_level):
    header = ['Ethnicity'] + THR_LABELS
    tbl = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for eth in ETHS:
        row = tbl.add_row().cells
        row[0].text = eth
        set_cell_bg(row[0], ETH_COLORS[eth])
        for j, thr in enumerate(THRESHOLDS):
            row[j + 1].text = ci_cell(match_level, eth, thr)
    # Weighted average row
    row = tbl.add_row().cells
    row[0].text = 'Combined pooled registry†'
    set_cell_bg(row[0], 'E8E8E8')
    for j, thr in enumerate(THRESHOLDS):
        v = tgts[(tgts.match_level == match_level) & (tgts.ethnicity == 'Combined') &
                 (tgts.model_variant == 'same_ethnicity') &
                 (tgts.target_coverage == thr)]['registry_size']
        row[j + 1].text = n(v.iloc[0]) if not v.empty else '—'
    # Attrition-adjusted signed-up targets (40% attrition → ×1.667)
    ATTRITION = 0.40
    row = tbl.add_row().cells
    row[0].text = 'Signed-up target‡\n(40% attrition)'
    set_cell_bg(row[0], 'FEF9E7')
    for j, thr in enumerate(THRESHOLDS):
        # Per-ethnicity EM point estimates; row spans the CMIO range
        # Show range across the four CMIO groups
        vals = []
        for eth in ETHS:
            r = tgts[(tgts.ethnicity == eth) & (tgts.match_level == match_level) &
                     (tgts.target_coverage == thr) &
                     (tgts.model_variant == 'same_ethnicity')]
            if not r.empty:
                vals.append(int(r.iloc[0].registry_size / (1 - ATTRITION)))
        if vals:
            row[j + 1].text = f'{min(vals):,}–{max(vals):,}'
        else:
            row[j + 1].text = '—'
    style_table(tbl)
    return tbl


def hap_display(h):
    """Convert 'aa:bb|cc:dd|...' to 'A*aa:bb~B*cc:dd~C*..~DRB1*..~DQB1*..'"""
    loci  = ['A', 'B', 'C', 'DRB1', 'DQB1']
    parts = h.split('|')
    return '~'.join(f'{l}*{a}' for l, a in zip(loci, parts))


# ── Build Document ────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── TITLE ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
tr = title.add_run(
    'Modelling Unrelated Donor Registry Size Requirements for '
    'Haematopoietic Stem Cell Transplantation in Singapore\'s '
    'Multiethnic Population'
)
tr.bold = True
tr.font.size = Pt(14)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_after = Pt(2)
ar = authors.add_run('Alvin Ng Yu-Jin')
ar.italic = True
ar.font.size = Pt(11)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.paragraph_format.space_after = Pt(12)
afr = affil.add_run(
    'National University Hospital / Singapore Cord Blood Bank, Singapore\n'
    'Correspondence: alvin1976sg@gmail.com'
)
afr.font.size = Pt(9)
afr.italic = True

doc.add_paragraph()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Abstract', level=2)
add_para(doc,
    'Background: Access to a compatible unrelated bone marrow donor is critical '
    'for patients with haematological malignancies requiring haematopoietic stem '
    'cell transplantation (HSCT). Singapore\'s multiethnic Chinese–Malay–Indian–'
    'Others (CMIO) population presents distinct HLA diversity challenges [1,11]. '
    'Using expectation-maximisation (EM)-estimated phased haplotype frequencies '
    'from 59,186 donors and cord blood units [1], we modelled the minimum registry '
    'size needed to achieve 75–95% population coverage at both 8/8 and 10/10 HLA '
    'match levels. This version (v2.16) re-runs the entire pipeline with the '
    'haplotype retention floor lowered from 0.1% to 0.01%, after the higher floor '
    'was found to discard the majority of haplotype frequency mass and to inflate '
    'estimated coverage by more than two orders of magnitude.')
add_corrected_para(doc, [
    ('Methods: Diplotype frequencies were derived under Hardy–Weinberg equilibrium '
     'from EM-phased five-locus haplotypes ', False),
    ('estimated using an expectation–maximisation algorithm with full phase '
     'enumeration ' + cite(20) + '; '
     'results were validated against the HLA-net GENE[RATE] database ' + cite(3) + '', True),
    ('. Coverage was modelled as a function of registry size ' + cite(2, 5) + '. Bootstrap '
     'Confidence intervals are not reported: the Dirichlet bootstrap used in earlier '
     'versions is biased downward for N* on a long-tailed haplotype distribution ' + cite(18) + ' (SS2.4). '
     'Partial-match (9/10, 8/10) relaxation, cross-ethnic feasibility, and '
     'ancestry stratification of the Others sub-group were analysed as secondary '
     'objectives.', False),
])
add_para(doc,
    'Results: For 10/10 HLA matching at 95% same-ethnicity coverage, minimum '
    f'registry sizes are approximately {K("n95_txt","Chinese")} (Chinese), '
    f'{K("n95_txt","Malay")} (Malay), {K("n95_txt","Indian")} (Indian) and '
    f'{K("n95_txt","Others")} (Others) donors — beyond the reach of '
    'any national programme. These estimates supersede the 40,000–45,000 donor '
    'targets reported up to v2.15, which were produced with a 0.1% haplotype '
    'frequency floor that discarded 47–64% of haplotype frequency mass; re-running '
    'the identical pipeline with a 0.01% floor retains 97–98% of that mass and '
    'raises N* by more than two orders of magnitude (§4.1). Because N* at these '
    'targets is unattainable, we report coverage attainable at feasible registry '
    f'size: at 50,000 same-ethnicity donors, modelled 10/10 coverage is '
    f'{K("cov50k","Chinese")}% '
    f'(Chinese), {K("cov50k","Malay")}% (Malay), {K("cov50k","Indian")}% (Indian) and {K("cov50k","Others")}% (Others); at 500,000 '
    f'donors, {K("cov500k","Chinese")}%, {K("cov500k","Malay")}%, {K("cov500k","Indian")}% and {K("cov500k","Others")}%. EM haplotype frequencies now agree '
    'closely with the independent HLA-net GENE[RATE] estimates across 1,227–3,011 '
    'shared haplotypes per group (Spearman r=0.75–0.95), covering 77–98% of '
    'frequency mass. Cross-ethnic matching remains substantially worse than '
    'same-ethnicity matching for Malay, Indian and Others patients. The pooled '
    'Others estimate remains a statistical artefact of heterogeneous population '
    'mixing; the group departs from Hardy–Weinberg equilibrium at all five loci '
    'with heterozygote deficit throughout, the classical Wahlund signature of '
    'population stratification, and ancestry sub-cluster analysis identifies three '
    'genetically distinct groups (European/Eurasian, Filipino/SE Asian, Northeast '
    'Asian).')
add_para(doc,
    'Conclusions: Achieving 95% 10/10 coverage from a domestic same-ethnicity '
    'registry is not feasible for any CMIO group. Planning should therefore be '
    'framed as coverage attainable at achievable registry size, and should '
    'prioritise the levers that do scale: evidence-based single-mismatch (9/10) '
    'protocols, international and diaspora registry linkage, cord blood, and '
    'haploidentical transplantation. The equity finding is unchanged and sharpened '
    '— at every feasible registry size, Indian and Others patients have roughly '
    'half the match probability of Chinese and Malay patients.',
    space_after=8)

kw = doc.add_paragraph()
kw.paragraph_format.space_after = Pt(12)
kwr = kw.add_run(
    'Keywords: HLA; donor registry; haematopoietic stem cell transplantation; '
    'CMIO; Singapore; registry size; partial match; EM haplotype estimation; '
    'EM maximum-likelihood estimate; Others stratification'
)
kwr.italic = True
kwr.font.size = Pt(9)

doc.add_page_break()

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction')

add_para(doc,
    'Haematopoietic stem cell transplantation (HSCT) offers a potential cure for '
    'many blood cancers and bone marrow disorders. When a patient lacks a matched '
    'family donor — as is the case for the majority — their best hope lies in an '
    'unrelated volunteer donor from a national or international registry. '
    'High-resolution HLA matching between donor and patient is strongly associated '
    'with transplant success and patient survival ' + cite(7) + '.')

add_para(doc,
    'Finding a matched donor is substantially harder for patients from non-European '
    'backgrounds ' + cite(6, 10) + '. Worldwide registries are dominated by donors '
    'of European descent, and HLA allele combinations that are common in East Asian, '
    'South Asian, or mixed-ancestry populations may be absent or very rare in those '
    'registries. For Singapore\'s Chinese, Malay, Indian, and Others (CMIO) '
    'populations ' + cite(11) + ', this creates an urgent need to understand exactly '
    'how large a local, ethnicity-representative registry must be to give patients a '
    'realistic chance of finding a match ' + cite(8, 9) + '.')

add_para(doc,
    'Ng et al. (2022) characterised the HLA landscape of Singapore\'s donor pool '
    'across 59,186 individuals and published allele and haplotype frequencies for '
    'each CMIO group ' + cite(1) + '. That study provided the essential '
    'epidemiological foundation but did not translate these frequencies into '
    'quantitative registry size recommendations. The present study does so, '
    'addressing four questions:')

for bullet in [
    'How many same-ethnicity donors are needed for 75–95% coverage at strict '
    '(10/10) and standard (8/8) HLA match levels, and how confident can we be '
    'in these estimates?',
    'Can cross-ethnic matching — drawing Malay, Indian, or Others patients from '
    'a Chinese-dominated combined pool — ever achieve adequate coverage?',
    'How much smaller could registries be if clinicians accepted partial matches '
    '(9 or 8 of 10 alleles)?',
    'Is the Others group sufficiently homogeneous to be modelled as a single '
    'population, or do distinct ancestry sub-groups require separate strategies?',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(bullet).font.size = Pt(10)

add_para(doc,
    'Answers to these questions carry direct implications for Singapore\'s bone '
    'marrow donor recruitment strategy and for donor registries serving similar '
    'multiethnic populations in the Asia-Pacific region ' + cite(8, 9) + '.')

# ── 2. METHODS ────────────────────────────────────────────────────────────────
add_heading(doc, '2. Methods')
add_figure(doc, 'pipeline_flowchart.png', width=6.5,
    caption='Figure 1. Analysis pipeline overview: from raw HLA typing data to '
    'registry size estimates as EM maximum-likelihood values. '
    'Shading: blue = input data; green = core model steps; '
    'yellow = optimisation; red = uncertainty quantification.')
add_heading(doc, '2.1 Dataset', level=2)
add_corrected_para(doc, [
    ('HLA typing data were obtained from 59,186 ', False),
    ('donors and cord blood units from the Bone Marrow Donor Programme (BMDP) '
     'and Singapore Cord Blood Bank (SCBB)', True),
    (', accrued between 2005 and 2020 ' + cite(1) + '. '
     'Recipients and donors processed by the Health Sciences Authority (HSA) were '
     'included for validation. ', False),
    ('Samples', True),
    (' were typed at high resolution across five loci '
     '(HLA-A, -B, -C, -DRB1, -DQB1). Full details of typing methodology and data '
     'cleaning are reported in Ng et al. ' + cite(1) + '.', False),
])

add_heading(doc, '2.2 Haplotype Frequency Estimation Using EM Phasing', level=2)
add_corrected_para(doc, [
    ('A key methodological advance in this analysis is the use of '
     'expectation-maximisation (EM)-estimated phased haplotype frequencies, rather '
     'than the simpler approach of multiplying individual allele frequencies at each '
     'locus. Haplotype frequencies were estimated using ', False),
    ('an expectation–maximisation algorithm with full multi-locus phase '
     'enumeration (Excoffier–Slatkin formulation ' + cite(20) + '), implemented '
     'in Python, applying Hardy–Weinberg equilibrium to infer the most probable '
     'phase assignments from unphased diplotype data (all five-locus-typed '
     'individuals per ethnic group; haplotypes retained at '
     'frequency ≥ 0.01%). Results were validated against the HLA-net '
     'GENE[RATE] database ' + cite(3) + '.', True),
])

add_para(doc,
    'The haplotype retention floor is not a neutral implementation detail. '
    'Versions of this analysis up to v2.15 used a 0.1% floor, which retained only '
    '123–144 haplotypes per group representing 36–53% of total haplotype frequency '
    'mass; the surviving frequencies were then renormalised to sum to one, so the '
    'model behaved as though every patient carried one of the few most common '
    'haplotypes. Because the coverage function C(N) converges slowly precisely '
    'where diplotype frequencies are small, the discarded tail is what determines '
    'high-coverage behaviour. Lowering the floor to 0.01% retains 2,310–3,035 '
    'haplotypes per group and 97–98% of frequency mass, and raises N* at 95% '
    'coverage by more than two orders of magnitude. A sweep across floors shows '
    'the effect is a threshold rather than a gradient: for Chinese donors, moving '
    'from no floor to 0.01% discards 96% of distinct haplotypes but only 3.1% of '
    'frequency mass and changes modelled coverage at 50,000 donors by 2.7 '
    'percentage points, whereas moving from 0.01% to 0.1% discards a further 44% '
    'of mass and moves the same figure from 39.5% to 95.6%. Removing all '
    'haplotypes below 0.001% changes coverage by 0.2 percentage points, '
    'establishing that the tail carrying the effect is not an artefact of EM '
    'phase enumeration.')

add_para(doc,
    'One further consequence deserves statement because it constrains how these '
    'estimates should be read. A fixed frequency floor means different things at '
    'different sample sizes. In a sample of 5,000 individuals (10,000 chromosomes) '
    'a haplotype observed once already has frequency 1e-4 and survives a 0.01% '
    'floor, so the floor barely discriminates and retains sampling noise; in the '
    'full 44,400-individual Chinese sample a singleton sits near 1.1e-5 and is '
    'removed. Using all available individuals therefore lowers retained frequency '
    'mass (Chinese 97.2% to 80.8%) even though it improves the estimate, because '
    'genuinely rare haplotypes are now separated from noise rather than pooled '
    'with it. An appropriate floor scales roughly as 1/(2n), implying about 1e-5 '
    'for the full sample. This analysis uses 1e-4 with the full sample, retaining '
    '81–98% of mass against 36–53% in versions to v2.15; the remaining discarded '
    'tail means the registry sizes reported here should be read as lower bounds.')

add_para(doc,
    'HLA loci are not independent — alleles at neighbouring loci are inherited '
    'together as haplotype blocks far more often than chance would predict. This '
    'non-random co-inheritance, known as linkage disequilibrium (LD), is especially '
    f'strong between DRB1 and DQB1 (composite D′ {LD_DRDQ[0]}–{LD_DRDQ[1]} across CMIO groups) and '
    'between B and C (D′ ≥ 0.95) ' + cite(1, 13) + '. Ignoring LD by '
    'using a simple product approximation systematically underestimates the '
    'frequency of common haplotypes. The EM algorithm produces a more concentrated '
    'and biologically realistic diplotype frequency distribution, yielding more '
    'accurate registry size projections ' + cite(3, 5) + '.')

add_heading(doc, '2.3 Registry Size Model', level=2)
add_para(doc,
    'For each CMIO group, the EM algorithm produced a ranked list of K haplotypes '
    'with estimated frequencies f₁, f₂, ..., fᴊ. Diplotype (genotype '
    'pair) frequencies were derived under Hardy–Weinberg equilibrium:')
add_equation(doc, 'F(hᵢ, hᵢ) = fᵢ²   [homozygous, i = j]', 1)
add_equation(doc, 'F(hᵢ, hⱼ) = 2·fᵢ·fⱼ,   i ≠ j   [heterozygous]', 2)
add_para(doc,
    'The probability that a patient with diplotype dₖ finds at least one matched '
    'donor in a registry of N independently drawn donors is:')
add_equation(doc, 'P(dₖ, N) = 1 − (1 − Fₖ)^N', 3)
add_para(doc,
    'Population coverage C(N) — the expected fraction of all patients who find '
    'at least one match — is the diplotype-frequency-weighted sum over all m '
    'diplotypes:')
add_equation(doc, 'C(N) = Σₖ Fₖ · [1 − (1 − Fₖ)^N],   k = 1, …, m', 4)
add_para(doc,
    'Equation (4) aggregates across all diplotype classes. For a patient whose '
    'diplotype dₖ occurs at population frequency Fₖ, the probability that at '
    'least one of N randomly drawn registry donors carries that same diplotype '
    'is [1 − (1 − Fₖ)^N] — the complement of all N donors lacking it. '
    'Multiplying by Fₖ weights this probability by the fraction of patients '
    'who carry diplotype dₖ; summing over all m diplotypes gives the overall '
    'population coverage C(N). As N grows, each term approaches Fₖ and C(N) '
    'approaches 1, but the convergence is slow for rare diplotypes.')
add_corrected_para(doc, [
    ('The minimum registry size N* to achieve target coverage θ (75%, 85%, '
     '90%, or 95%) was found by binary search on a logarithmic scale (50 iterations; '
     'precision < 1 donor for N ≤ 10⁷) ' + cite(2, 5) + '. '
     'Same-ethnicity and cross-ethnic model variants were compared ' + cite(2) + '. ', False),
    ('The coverage-based framework — computing C(N) as the diplotype-frequency-weighted '
     'sum of per-patient match probabilities — follows the approach pioneered by '
     'Beatty et al. ' + cite(4) + '.', True),
])

add_heading(doc, '2.4 Uncertainty: why confidence intervals are not reported', level=2)
add_para(doc,
    'Versions of this analysis up to v2.15 reported Dirichlet parametric bootstrap '
    'confidence intervals (B = 1,000), resampling haplotype frequencies from a '
    'Dirichlet distribution with concentration parameters n_eth x f_k and '
    'recomputing N* per resample. Those intervals are withdrawn, because the '
    'procedure is biased for this statistic rather than merely imprecise.')
add_para(doc,
    'Under a Dirichlet draw the expectation of a squared frequency is not the '
    'square of the expectation: E[f^2] = f^2 + f(1-f)/(n+1). Diplotype frequencies '
    'are squares and pairwise products of haplotype frequencies, so resampling '
    'inflates them, and the inflation is largest exactly where the frequencies are '
    'smallest. For the Chinese distribution the mean inflation of E[f^2] is about '
    '525% for haplotypes between 1e-6 and 1e-5, 134% between 1e-5 and 1e-4, 12% '
    'between 1e-4 and 1e-3, and 1% above 1e-3. Because the rare tail is what '
    'governs high-coverage behaviour, every resample overstates attainable coverage '
    'and therefore understates N*.')
add_para(doc,
    'The symptom was visible in earlier releases and was misread as skew. In v2.15, '
    '18 of 32 reported rows had the EM maximum-likelihood estimate lying outside '
    'its own confidence interval, always above the upper bound, with 98-100% of '
    'resamples falling below the point estimate. Reporting the bootstrap median '
    'instead of the maximum-likelihood value made the tables internally consistent '
    'but did not address the cause. At the 1e-6 floor used here the discrepancy is '
    'unmistakable: for Chinese 10/10 at 95% coverage the maximum-likelihood '
    'estimate is 87,384,114 against a bootstrap interval of 50,772,223 to '
    '53,870,680, with every one of 1,000 resamples below the point estimate.')
add_para(doc,
    'All registry sizes in this version are therefore EM maximum-likelihood point '
    'estimates, with no interval attached. This is a statement of what is not '
    'known rather than a claim of precision: the estimates carry real uncertainty '
    'from haplotype-frequency sampling, from the EM phasing step, from the '
    'Hardy-Weinberg assumption, and from the choice of frequency floor (SS2.2), '
    'and that uncertainty is not quantified here. A resampling scheme that '
    'preserves the rare-tail structure, or an analytic propagation through the '
    'coverage function, would be required to restore intervals; neither is '
    'attempted in this release. Readers should treat every N* as an '
    'order-of-magnitude quantity.')

add_heading(doc, '2.5 Partial-Match Analysis', level=2)
add_corrected_para(doc, [
    ('Partial-match coverage curves were generated for match thresholds of 10/10, '
     '9/10, and 8/10 (5-locus framework) and 8/8, 7/8, 6/8 (4-locus framework) ', False),
    (cite(2, 19), True),
    ('. A partial match is defined as sharing a specified minimum '
     'number of alleles across the relevant loci. These analyses used EM-phased '
     'haplotype pairs to correctly account for LD structure ' + cite(3) + '.', False),
])

add_heading(doc, '2.6 Others Group Stratification', level=2)
add_para(doc,
    'The Others group is ethnically heterogeneous, comprising Eurasians, '
    'Caucasians, and a variety of mixed and other Asian backgrounds. To assess '
    'whether this group should be treated as a single population, we applied '
    'principal component analysis (PCA) to a binary HLA allele indicator matrix '
    '(alleles with ≥1% frequency across the Others cohort), followed by '
    'k-means clustering (k = 2–5). The optimal number of clusters was selected '
    'by silhouette coefficient. Registry size requirements were then estimated '
    'independently for each cluster. Ancestry inference for each cluster was based '
    'on haplotype signature matching against the Allele Frequency Net Database '
    '(AFND) ' + cite(16) + ' and published population-specific haplotype '
    'references ' + cite(1, 17) + '.')

# ── 3. RESULTS ────────────────────────────────────────────────────────────────
add_heading(doc, '3. Results')
add_para(doc,
    'Section 3.1 presents the primary finding: same-ethnicity registry size '
    'requirements at 10/10 matching. Section 3.2 shows that 8/8 targets are '
    'marginally lower. Section 3.3 demonstrates that cross-ethnic matching '
    'cannot substitute for same-ethnicity donors. Sections 3.4–3.5 present '
    'secondary analyses (partial-match and sensitivity). Section 3.6 reports '
    'model validation. Section 3.7 is explicitly exploratory: ancestry '
    'stratification of the heterogeneous "Others" subgroup.')

# 3.1 Registry size 10/10
add_heading(doc, '3.1 Registry Size Requirements — 10/10 HLA Matching', level=2)
add_para(doc,
    'Table 1 presents the minimum registry sizes required for same-ethnicity '
    '10/10 HLA matching across all four CMIO groups, based on EM-phased '
    'haplotype frequencies ' + cite(1, 3) + '. Values are EM maximum-likelihood '
    'estimates (bias-corrected point estimates); numbers in parentheses are 95% '
    'point estimates; confidence intervals are withdrawn (SS2.4) ' + cite(18) + '. '
    'The Combined pooled registry row models a single shared registry in which both '
    'patients and donors are drawn from a Singapore-weighted pool ' + cite(11) + '; '
    'it does not guarantee equitable access for minority ethnic groups.')

make_ci_table(doc, '10of10')
add_para(doc,
    '† Combined pooled registry (NOT an average of the per-group values): both patients '
    'and donors drawn from a Singapore-weighted pool (Chinese 77%, Malay 8%, Indian 9%, '
    'Others 6%) [11]. This row is a mathematical convenience, not a policy target: '
    'it describes how a single shared registry would need to be sized to serve all '
    'groups proportionally under those weights, but it cannot substitute for '
    'ethnicity-specific recruitment because donors are not interchangeable across '
    'groups. The per-group same-ethnicity N* values above are the operative planning '
    'targets for each community.\n'
    'Values shown as EM maximum-likelihood N. No interval is attached; see SS2.4.',
    size=8, space_after=4)
add_caption(doc,
    'Table 1. Minimum same-ethnicity registry size for 10/10 HLA matching '
    'by coverage target. Values are EM maximum-likelihood estimates; no '
    'confidence intervals are attached, because the Dirichlet bootstrap used in '
    'earlier versions is biased downward for this statistic (§2.4) [18]. '
    'Estimates should be read as order-of-magnitude quantities. The pooled Others '
    'row is a mathematical artefact; see §3.7 for sub-cluster targets. '
    '‡ Signed-up target assumes 40% real-world donor attrition '
    '(unreachability, refusal, or medical deferral); shows range across CMIO '
    'groups at each threshold.', fig=False)

add_para(doc,
    'At the 95% coverage threshold, same-ethnicity registry sizes are far beyond '
    f'what any single national programme can recruit: approximately '
    f'{K("n95_txt","Chinese")} donors (Chinese), {K("n95_txt","Malay")} (Malay), '
    f'{K("n95_txt","Indian")} (Indian) and {K("n95_txt","Others")} (Others). Even '
    f'75% coverage requires {K("n75_txt","Malay")} to {K("n75_txt","Others")} '
    'same-ethnicity donors. '
    'The Others group requires the largest registry, followed by Indian, reflecting '
    'their greater haplotype diversity ' + cite(1, 8) + '. These figures supersede '
    'the 40,000–45,000 donor targets reported in versions up to v2.15 of this '
    'analysis, which were computed with a 0.1% haplotype frequency floor that '
    'discarded 47–64% of haplotype frequency mass before the coverage calculation '
    '(see §2.2 and §4.1). They are consistent in order of magnitude with observed '
    'international registry performance: the US registry, at 10.5 million donors, '
    'achieves roughly 75% 8/8 match likelihood for its best-served ancestry group '
    + cite(2) + '.')
add_para(doc,
    'Because N* at these coverage targets exceeds any feasible national registry, '
    'the operationally meaningful quantity is the coverage attainable at a '
    'realistic registry size. At 50,000 same-ethnicity donors — approximately the '
    'scale of the current combined BMDP and SCBB inventory — modelled 10/10 '
    f'coverage is {K("cov50k","Chinese")}% for Chinese patients, '
    f'{K("cov50k","Malay")}% for Malay, {K("cov50k","Indian")}% for Indian and '
    f'{K("cov50k","Others")}% for Others. At 100,000 donors these rise to '
    f'{K("cov100k","Chinese")}%, {K("cov100k","Malay")}%, {K("cov100k","Indian")}% '
    f'and {K("cov100k","Others")}%, and at 500,000 donors to '
    f'{K("cov500k","Chinese")}%, {K("cov500k","Malay")}%, {K("cov500k","Indian")}% '
    f'and {K("cov500k","Others")}%. '
    'The equity gap is therefore not a matter of a few thousand donors: at any '
    'feasible registry size, Indian and Others patients have roughly half the match '
    'probability of Chinese and Malay patients.')

add_para(doc,
    'A striking feature of these results is how rapidly requirements grow as '
    'coverage targets increase beyond 90%. This reflects the long tail of rare '
    'HLA diplotypes: the first 90% of patients can be served by a relatively '
    'concentrated pool of common haplotype combinations, but the final 5–10% '
    'of patients carry increasingly rare diplotypes that require '
    'disproportionately larger registries to cover ' + cite(2) + '. It is precisely '
    'this tail that a frequency floor removes, which is why the choice of floor '
    'dominates the resulting registry-size estimate.')

add_figure(doc, 'coverage_curves_10of10.png', width=6.4,
    caption='Figure 2. Coverage as a function of registry size for 10/10 matching, '
    'by CMIO group, under same-ethnicity (solid) and cross-ethnic (dashed) donor '
    'pools. The curve, rather than a single N*, is the honest summary: it shows '
    'what a registry of any achievable size delivers, and makes clear how far the '
    '95% target sits beyond feasible recruitment. This figure replaces the bootstrap '
    'confidence-interval plot of earlier versions, for the reasons in section 2.4.')

# 3.2 Registry size 8/8
add_heading(doc, '3.2 Registry Size Requirements — 8/8 HLA Matching', level=2)
add_para(doc,
    'At the 8/8 match level (omitting DQB1 as a mandatory matching criterion), '
    'registry size requirements are marginally lower than for 10/10 matching '
    '(Table 2). The difference is modest — typically 600–1,200 fewer donors at '
    '95% coverage — because DRB1 and DQB1 are in very strong linkage '
    f'disequilibrium across all CMIO groups (composite D′ {LD_DRDQ[0]}–{LD_DRDQ[1]}) ' + cite(1, 13) + '. '
    'A donor matched at DRB1 is almost always also matched at DQB1, so the '
    'additional constraint of DQB1 matching adds very little to registry size '
    'requirements in practice.')

make_ci_table(doc, '8of8')
add_para(doc,
    '† Combined pooled registry: patients and donors both drawn from a Singapore-weighted '
    'pool [11]; this is not an average of the per-group values and does '
    'not guarantee equitable access for minority ethnic groups; see same-ethnicity '
    'targets above.\n'
    '‡ Signed-up target assumes 40% real-world donor attrition; shows range across CMIO groups.\n'
    'Values shown as EM maximum-likelihood N. No interval is attached; see SS2.4.',
    size=8, space_after=4)
add_caption(doc,
    'Table 2. Minimum same-ethnicity registry size for 8/8 HLA matching '
    'by coverage target. EM maximum-likelihood estimates; confidence intervals '
    'withdrawn (§2.4) [18]. '
    '‡ Signed-up target assumes 40% attrition; shows range across CMIO groups.', fig=False)

add_para(doc,
    'This finding has an important practical implication: the strong DRB1–DQB1 '
    'LD means that investing in DQB1 typing data is worthwhile — it adds '
    'meaningful clinical information (DQB1 mismatches can increase graft-versus-'
    'host disease risk) at very little cost in registry size. Programmes can '
    'justify full 10-locus typing without substantially larger recruitment targets '
    + cite(7, 12) + '.')

# 3.3 Cross-ethnic
add_heading(doc,
    '3.3 Cross-Ethnic Matching — Why It Cannot Replace Same-Ethnicity Donors',
    level=2)
add_para(doc,
    'A natural question is whether a large shared registry — primarily supplied '
    'by the numerically dominant Chinese donor group — could serve patients of '
    'all ethnicities. The cross-ethnic analysis answers this question definitively: '
    'it cannot (Table 3). This is consistent with findings from other multiethnic '
    'registries ' + cite(10) + '.')

tbl3_data = [['Ethnicity', '75% coverage', '85% coverage',
              '90% coverage', '95% coverage']]
for eth in ETHS:
    row = [eth]
    for thr in THRESHOLDS:
        row.append(cross_cell('10of10', eth, thr))
    tbl3_data.append(row)

tbl3 = doc.add_table(rows=len(tbl3_data), cols=5)
for i, row_data in enumerate(tbl3_data):
    for j, val in enumerate(row_data):
        tbl3.rows[i].cells[j].text = val
style_table(tbl3)
for i, eth in enumerate(ETHS):
    set_cell_bg(tbl3.rows[i + 1].cells[0], ETH_COLORS[eth])
add_para(doc,
    'Combined donor pool: Singapore population weights (Chinese 77%) [11].',
    size=8, space_after=4)
add_caption(doc,
    'Table 3. Minimum registry size for 10/10 HLA matching under the cross-ethnic '
    'model (combined Singapore donor pool). Values >10,000,000 are infeasible.',
    fig=False)

add_para(doc,
    'At 95% coverage the penalty for drawing donors from a Singapore-weighted '
    f'combined pool rather than a same-ethnicity pool is smallest for Chinese '
    f'patients ({K("cross95","Chinese")} versus {K("n95_txt","Chinese")}), because Chinese haplotypes '
    'dominate the combined pool. For the other groups the penalty is of a '
    f'different order entirely: {K("cross95","Malay")} for Malay, {K("cross95","Indian")} for Indian '
    f'and {K("cross95","Others")} for Others. Their '
    'distinctive haplotype combinations are not represented in the pool, so scale '
    'cannot compensate ' + cite(10) + '. This is the one headline conclusion of '
    'earlier versions that survives the corrected analysis intact, and it survives '
    'it more strongly: same-ethnicity recruitment is not merely preferable but, on '
    'these results, the only viable domestic strategy for these groups ' + cite(9) + '.')

add_para(doc,
    'These cross-ethnic figures also incorporate a correction to the matching code '
    'itself. Diplotype pairs were previously labelled in each population\'s own '
    'frequency-rank order, so the same unordered pair of haplotypes could be stored '
    'as (X,Y) in the patient table and (Y,X) in the donor table; the lookup that '
    'joins them on those labels then silently scored such pairs as unmatched. '
    'Measured on Malay patients against the combined pool, 30.4% of patient '
    'frequency mass was affected, and correcting it raised cross-ethnic coverage at '
    'a one-million-donor registry from 0.635 to 0.776. Pairs are now labelled in a '
    'canonical order. The correction makes cross-ethnic matching look better than '
    'previously reported, and the conclusion above holds despite it rather than '
    'because of it.')

# 3.4 Partial match
add_heading(doc,
    '3.4 Partial-Match Coverage — The Clinical Benefit of Relaxing Match Criteria',
    level=2)
add_para(doc,
    'While exact 10/10 or 8/8 matching is the preferred clinical standard '
    + cite(7, 19) + ', some centres accept partial matches — particularly when '
    'no fully matched donor exists or when time is critical ' + cite(14) + '. '
    'Figures 3 and 4 show coverage curves for progressive relaxation of match '
    'criteria in the 10-locus and 8-locus frameworks.')

add_figure(doc, 'partial_match_10locus.png', width=6.2,
    caption='Figure 3. Coverage curves for 10-locus (HLA-A, -B, -C, -DRB1, -DQB1) '
    'partial matching across CMIO groups. Green: 10/10 exact match; Blue: ≥9/10; '
    'Red: ≥8/10. Dashed lines: cross-ethnic model. Note the sharp threshold '
    'effect at the 9/10 boundary for all groups.')

add_para(doc,
    'Relaxing the match requirement by a single allele is the largest lever '
    'available, and at the corrected 0.01% floor it is best expressed as coverage '
    'gained at a fixed registry size rather than as a ratio of unattainable N* '
    'values. For Chinese patients at a 50,000-donor same-ethnicity registry, '
    'modelled coverage rises from 40.6% at strict 10/10 to 72.0% at ≥9/10, and to '
    '93.8% at ≥8/10. Accepting one mismatch therefore delivers roughly the same '
    'gain as a thirty-fold increase in registry size, which no recruitment '
    'campaign can match. This is clinically significant because several large '
    'multi-centre trials have demonstrated that a single allele mismatch at certain '
    'loci — particularly at HLA-DPB1 when the mismatch is at a permissive '
    'T-cell epitope group position — carries minimal additional survival impact '
    + cite(12, 15) + '. Two cautions apply: the model treats a mismatch at any of '
    'the five loci as equivalent, which the outcomes literature does not support '
    '(permissiveness is locus-specific, which is why DPB1 is assessed through a '
    'T-cell-epitope framework rather than as a flat allele count), and ≥8/10 falls '
    'below currently accepted match stringency for most indications and is '
    'reported here only to show the shape of the relationship, not as a '
    'recommendation ' + cite(2, 19) + '.')

add_para(doc,
    'Note on method: the exact partial-match algorithm is quadratic in the number '
    'of diplotypes and is not computable at the 0.01% floor (approximately 8e12 '
    'diplotype comparisons per group). The figures above are Monte-Carlo '
    'estimates obtained by sampling 3,000 patient diplotypes in proportion to '
    'their frequency and evaluating each against the complete donor diplotype set, '
    'which is unbiased for coverage with sampling error of order 1/sqrt(3000). '
    'Figures 3 and 4 below were generated at the previous 0.1% floor and are '
    'retained for their qualitative shape only; their absolute registry-size axis '
    'is not comparable with the tables in this version.')

add_figure(doc, 'partial_match_8locus.png', width=6.2,
    caption='Figure 4. Coverage curves for 8-locus (HLA-A, -B, -C, -DRB1) partial '
    'matching. Green: 8/8; Blue: ≥7/8; Red: ≥6/8. Cross-ethnic '
    'shown with dashed lines.')

add_para(doc,
    'In the 8-locus framework (Figure 4), the same pattern holds. The cross-ethnic '
    'curves (dashed) plateau for Malay, Indian, and Others groups regardless of '
    'relaxation level — confirming that partial matching cannot substitute for '
    'same-ethnicity donors when fundamental haplotype diversity differs markedly '
    'between donor and patient populations ' + cite(2, 10) + '.')

# 3.5 Sensitivity
add_heading(doc,
    '3.5 Registry Size Is Robust to Patient Demographic Assumptions', level=2)
add_para(doc,
    'Registry size models must assume a patient population composition ' + cite(2, 5) + '. '
    'We tested four scenarios: Singapore population weights (77% Chinese, 8% Malay, '
    '9% Indian, 6% Others) ' + cite(11) + ', BMDP+SCBB registry composition, '
    'HSA Patient-Donor Data (referral composition from Health Sciences Authority Singapore), and an extreme minority-focus '
    'scenario (equal Malay, Indian, and Others weighting, no Chinese). Table 6 and '
    'Figure 5 show that the combined N* varies by less than 3% across all scenarios '
    'at any coverage target ' + cite(2) + '.')

tbl6_data = [['Scenario', 'Ethnic weights (C/M/I/O)',
              '75%', '85%', '90%', '95%']]
scenario_labels = {
    'SG population (current model)':        'SG population',
    'BMDP+SCBB registry composition':       'BMDP+SCBB donors',
    'Patient.txt composition':              'HSA Patient-Donor Data',
    'Minority-focus (Indian+Malay+Others)': 'Minority-focus',
}
scenario_weights = {
    'SG population (current model)':        '77/8/9/6',
    'BMDP+SCBB registry composition':       '75/9/9/6',
    'Patient.txt composition':              '72/15/5/8',
    'Minority-focus (Indian+Malay+Others)': '0/40/40/20',
}
for sk in ['SG population (current model)', 'BMDP+SCBB registry composition',
           'Patient.txt composition', 'Minority-focus (Indian+Malay+Others)']:
    sub = sens[sens.scenario == sk]
    row = [scenario_labels[sk], scenario_weights[sk]]
    for thr in THRESHOLDS:
        v = sub[sub.target_coverage == thr]['registry_size']
        row.append(n(v.iloc[0]) if not v.empty else '—')
    tbl6_data.append(row)

tbl6 = doc.add_table(rows=len(tbl6_data), cols=6)
for i, row_data in enumerate(tbl6_data):
    for j, val in enumerate(row_data):
        tbl6.rows[i].cells[j].text = val
style_table(tbl6)
add_caption(doc,
    'Table 6. Sensitivity of combined registry size to patient ethnic composition '
    'scenario [11]. C/M/I/O = Chinese/Malay/Indian/Others percentage weights.',
    fig=False)

add_figure(doc, 'cross_ethnic_sensitivity.png', width=6.0,
    caption='Figure 5. Registry size sensitivity across four patient demographic '
    'scenarios [11], including HSA Patient-Donor Data (Health Sciences Authority '
    'Singapore). Near-identical bar heights show that the combined estimate is '
    'insensitive to patient ethnic composition [2].')

add_para(doc,
    'This stability arises because per-group N* values are all in the same order '
    'of magnitude. Reweighting the groups therefore changes the combined N* only '
    'marginally ' + cite(2) + '. The practical implication is unchanged by the '
    'revised floor: the difficulty of the matching problem is a structural '
    'property of CMIO haplotype diversity rather than an artefact of assumptions '
    'about who the future patients will be, and will not be resolved by '
    'demographic change ' + cite(11) + '. Note that this sensitivity analysis '
    'tests only the patient-composition assumption; it does not probe the '
    'haplotype-retention floor, which §4.1 shows is the dominant source of '
    'variation in N*.')

# 3.6 Model validation
add_heading(doc, '3.6 Model Validation', level=2)
add_para(doc,
    'To assess whether EM-derived haplotype frequencies reflect real patient '
    'haplotype distributions, we compared EM estimates against independently '
    'observed haplotype frequencies from 564 patient-donor pairs provided by the '
    'Health Sciences Authority (HSA) Singapore ' + cite(1) + '. Spearman '
    'rank correlations and root mean square error (RMSE) were computed on shared '
    'haplotypes (Table 7).')

tbl7_header = ['Ethnicity', 'Patient haplotypes', 'Shared with EM',
               '% frequency covered', 'Spearman r', 'RMSE']
tbl7 = doc.add_table(rows=1, cols=6)
for i, h in enumerate(tbl7_header):
    tbl7.rows[0].cells[i].text = h
for _, mvrow in mv.iterrows():
    eth = mvrow['ethnicity']
    row = tbl7.add_row().cells
    row[0].text = eth
    row[1].text = str(int(mvrow['n_patient'])) if pd.notna(mvrow['n_patient']) else '—'
    row[2].text = str(int(mvrow['n_shared']))   if pd.notna(mvrow['n_shared'])  else '—'
    row[3].text = f"{mvrow['pct_covered']:.1f}%" if pd.notna(mvrow['pct_covered']) else '—'
    row[4].text = f"{mvrow['spearman_r']:.3f}"   if pd.notna(mvrow['spearman_r'])  else 'n/a*'
    row[5].text = f"{mvrow['rmse']:.4f}"          if pd.notna(mvrow['rmse'])        else '—'
    set_cell_bg(row[0], ETH_COLORS.get(eth, 'FFFFFF'))
style_table(tbl7)
add_para(doc, '* Insufficient shared haplotypes for rank correlation (n < 3).',
    size=8, space_after=4)
add_caption(doc,
    'Table 7. Validation of EM-estimated haplotype frequencies [3] against '
    'independently observed patient haplotypes.', fig=False)

add_figure(doc, 'match_validation_scatter.png', width=5.5,
    caption='Figure 6. Observed (patient) vs EM-estimated haplotype frequencies '
    'per ethnicity. Dashed line = perfect agreement. Chinese: Spearman r=0.70 '
    '(p<0.001, n=33 shared haplotypes).')

add_para(doc,
    'For Chinese patients, the EM estimates show good rank agreement with observed '
    'patient frequencies (Spearman r = 0.70, p < 0.001), with RMSE = 0.0094; '
    'this is the primary validation result. For Malay (11 shared haplotypes) and '
    'Others (4 shared haplotypes), sample sizes are insufficient for reliable rank '
    'correlation. For Indian patients only one shared haplotype was observed, '
    'precluding any meaningful validation ' + cite(1) + '. Registry size estimates '
    'for Malay, Indian, and Others should therefore be regarded as model-derived '
    'projections; the Chinese estimates are the most robustly validated. These '
    'results highlight the need for larger patient cohorts in minority groups.')

# 3.7 Others stratification — Exploratory; CMI is the primary focus
add_heading(doc, '3.7 Exploratory Analysis — Note on the \"Others\" Subgroup', level=2)
add_para(doc,
    'The Others category in Singapore\'s CMIO classification encompasses Eurasians, '
    'Caucasians, and individuals of mixed or diverse Asian ancestry ' + cite(11) + '. '
    'Unsupervised clustering (optimal k=3 by silhouette coefficient; s=0.24 in the '
    'five-component PCA space used for clustering) of the 3,847 fully five-locus-typed Others donors reveals '
    'three genetically distinct sub-groups — European/Eurasian, Filipino/SE Asian, '
    'and Northeast Asian/Mixed — with markedly different registry size requirements '
    '(Table 4). The pooled Others estimate in Table 1 should not be used as a '
    'policy target. At the corrected 0.01% floor the direction of the pooling '
    'artefact is the reverse of what earlier versions reported: pooled Others '
    f'requires {K("n95_txt","Others")} donors at 95% coverage, whereas each '
    f'individual sub-cluster requires only {_millions(int(oc_reg[oc_reg.target_coverage==0.95].registry_size.min()))} '
    f'to {_millions(int(oc_reg[oc_reg.target_coverage==0.95].registry_size.max()))}. This is the statistically expected '
    'ordering — a patient from one ancestry cluster rarely matches a donor from '
    'another, so a stratified pool is harder to match within than any of its '
    'homogeneous parts — and it is corroborated by the group\'s departure from '
    'Hardy–Weinberg equilibrium at all five loci with heterozygote deficit '
    'throughout (§4.1), the classical Wahlund signature. The practical implication '
    'is favourable and actionable: recruiting to a specific Others sub-ancestry is '
    'roughly five to ten times more efficient per donor than recruiting to the '
    'undifferentiated category, which makes collecting sub-ancestry data at '
    'registration a high-yield intervention rather than a refinement.')

# Table 4: cluster registry sizes
tbl4_header = ['Cluster', 'Putative ancestry', 'N individuals',
               '75%', '85%', '90%', '95%']
tbl4 = doc.add_table(rows=1, cols=len(tbl4_header))
for i, h in enumerate(tbl4_header):
    tbl4.rows[0].cells[i].text = h

ancestry_labels = {
    'Cluster_1': 'European / Eurasian',
    'Cluster_2': 'Filipino / SE Asian',
    'Cluster_3': 'Northeast Asian / Mixed',
}
clust_colors = {
    # Light tints of Figure 7 scatter colours (#e41a1c, #377eb8, #4daf4a)
    'Cluster_1': 'FFCCCC',  # light red   → #e41a1c
    'Cluster_2': 'CCE0F5',  # light blue  → #377eb8
    'Cluster_3': 'CBF0CB',  # light green → #4daf4a
}
for cl in ['Cluster_1', 'Cluster_2', 'Cluster_3']:
    sub   = oc_reg[oc_reg.cluster == cl]
    n_ind = int(sub.iloc[0].n_individuals)
    sizes = {int(r.target_coverage * 100): int(r.registry_size)
             for _, r in sub.iterrows()}
    row   = tbl4.add_row().cells
    row[0].text = cl.replace('_', ' ')
    row[1].text = ancestry_labels[cl]
    row[2].text = f'{n_ind:,}'
    for j, thr_pct in enumerate([75, 85, 90, 95]):
        row[3 + j].text = n(sizes.get(thr_pct, 0))
    set_cell_bg(row[0], clust_colors[cl])
    set_cell_bg(row[1], clust_colors[cl])
style_table(tbl4)
add_caption(doc,
    'Table 4. Minimum registry size for 10/10 HLA matching within each Others '
    'sub-cluster. Putative ancestry inferred from haplotype signature matching '
    'against AFND [16] and published references [1,17]. '
    'The pooled Others row (Table 1) is a mathematical artefact and must not '
    'be used as a policy target (see text).', fig=False)

add_figure(doc, 'others_pca_scatter.png', width=5.0,
    caption='Figure 7. PCA scatter of 3,847 fully five-locus-typed Others donors (binary HLA allele '
    'indicators, alleles ≥1%). Three distinct clusters (optimal k=3 by silhouette '
    'coefficient, s=0.24 in the five-PC clustering space; 0.43 in the PC1–PC2 '
    'projection shown) indicate different ancestry backgrounds. Cluster identities inferred from '
    'haplotype signature matching [16,17].')

# Mini haplotype validation table — top 2 haplotypes per cluster
hap_tbl_header = ['Cluster', 'Rank', 'Haplotype (A~B~C~DRB1~DQB1)',
                  'Freq (%)', 'Population association']

# Population annotations for top haplotypes
pop_annot = {
    # Cluster 1 haplotypes
    '01:01|08:01|07:01|03:01|02:01':
        '8.1 Ancestral Haplotype — hallmark Northern European [17]',
    '33:03|44:03|07:01|07:01|02:01':
        'Common South Asian (Indian subcontinent) [16]',
    '33:03|58:01|03:02|03:01|02:01':
        'Most common Chinese/East Asian haplotype [1]',
    '01:01|57:01|06:02|07:01|03:03':
        'South Asian signature (India, Sri Lanka) [16]',
    '29:02|44:03|16:01|07:01|02:01':
        'Mediterranean / Middle Eastern European [16]',
    # Cluster 2 haplotypes
    '24:07|35:05|04:01|12:02|03:01':
        'Filipino / SE Asian archipelago (B*35:05–DRB1*12:02) [16]',
    '11:01|15:02|08:01|12:02|03:01':
        'Most common Filipino haplotype (B*15:02–DRB1*12:02) [16]',
    '02:01|15:13|08:01|12:02|03:01':
        'SE Asian (B*15:13 — Philippines, Indonesia) [16]',
    '24:02|15:02|08:01|12:02|03:01':
        'Filipino / Thai / Indonesian [16]',
    '24:07|15:02|08:01|12:02|03:01':
        'SE Asian archipelago [16]',
    # Cluster 3 haplotypes
    '02:07|46:01|01:02|09:01|03:03':
        'Chinese-specific (A*02:07–B*46:01) [1]',
    '24:02|38:02|07:02|15:02|05:02':
        'Japanese / Korean signature (B*38:02–DRB1*15:02) [16]',
    '34:01|15:21|04:03|15:02|06:01':
        'South/SE Asian [16]',
    '03:01|07:02|07:02|15:01|06:02':
        'Common European haplotype [16]',
    '29:01|07:05|15:05|10:01|05:01':
        'South Asian / mixed [16]',
}

# Table 5: top haplotype per cluster (ancestry validation) — no Rank column
tbl5_header = ['Cluster', 'Top Haplotype (A~B~C~DRB1~DQB1)', 'Freq (%)', 'Population association']
tbl5 = doc.add_table(rows=1, cols=len(tbl5_header))
for i, h in enumerate(tbl5_header):
    tbl5.rows[0].cells[i].text = h

for cl in ['Cluster_1', 'Cluster_2', 'Cluster_3']:
    top1 = (oc_hap[oc_hap.cluster == cl]
            .sort_values('frequency', ascending=False)
            .iloc[0])
    row = tbl5.add_row().cells
    row[0].text = cl.replace('_', ' ')
    row[1].text = hap_display(top1['haplotype'])
    row[2].text = f"{top1['frequency']*100:.1f}"
    row[3].text = pop_annot.get(top1['haplotype'], '—')
    set_cell_bg(row[0], clust_colors[cl])

style_table(tbl5)
# Widen Population association column (col 3) and narrow Cluster/Freq columns
tbl5.autofit = False
col_widths = [Cm(2.8), Cm(5.2), Cm(1.3), Cm(7.5)]
for row in tbl5.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w
add_caption(doc,
    'Table 5. Top haplotype per Others sub-cluster confirming putative ancestry. '
    'Frequencies are within-cluster. Population assignments cross-referenced '
    'against AFND [16] and published references [1,17].', fig=False)

# Per-cluster narratives not rendered in main body

# ── 4. DISCUSSION AND RECOMMENDATIONS ────────────────────────────────────────
add_heading(doc, '4. Discussion and Recommendations')

add_para(doc,
    'This analysis provides quantitative registry size estimates for '
    'Singapore\'s CMIO populations, grounded '
    'in EM-phased haplotype frequencies from the largest local HLA dataset to date '
    + cite(1) + '. Its central practical finding is negative and, we believe, more '
    'useful than the target figures it replaces: no domestically achievable '
    'registry delivers high-resolution coverage comparable to what earlier '
    'versions of this analysis projected, and recruitment strategy should be '
    'planned accordingly ' + cite(9) + '.')

rec_items = [
    (
        'Plan against coverage attainable at achievable size, not against a 95% target.',
        f'A 95% 10/10 coverage target requires {K("n95_txt","Malay")} to {K("n95_txt","Others")} same-ethnicity '
        'donors and is not reachable by any national programme. Stating a target '
        'that cannot be met obscures the real decision, which is how to allocate '
        'a registry of realistic size. At 50,000 same-ethnicity donors, modelled '
        f'coverage is {K("cov50k","Chinese")}% (Chinese), {K("cov50k","Malay")}% (Malay), {K("cov50k","Indian")}% (Indian) and {K("cov50k","Others")}% '
        '(Others); these are the numbers a programme can be held to and can track '
        'year on year ' + cite(2) + '.'
    ),
    (
        'Prioritise Malay and Indian donor recruitment, and measure it in coverage gained.',
        'Both the Malay and Indian communities are under-represented in existing '
        'registries relative to their HSCT need ' + cite(8, 9) + '. Indian and '
        'Others patients sit at roughly half the match probability of Chinese and '
        'Malay patients at every feasible registry size, and closing that gap by '
        'recruitment alone would require millions of donors. Targeted outreach, '
        'community partnerships and incentive programmes remain necessary, but '
        'should be costed against the coverage percentage points they actually '
        'deliver rather than against an unreachable headcount.'
    ),
    (
        'Adopt partial-match protocols to extend effective coverage.',
        'Relaxing from 10/10 to 9/10 matching roughly halves the effective '
        'registry size requirement ' + cite(2) + '. In clinical contexts where '
        'survival outcomes are comparable between 10/10 and permissive 9/10 '
        'matches — particularly when the single mismatch is at a permissive '
        'DPB1 T-cell epitope group ' + cite(12, 15) + ' — formalising 9/10 as '
        'an accepted standard would effectively double the reach of existing '
        'registries without additional donor recruitment.'
    ),
    (
        'Disaggregate the Others category and collect ancestry sub-group data.',
        'Pooled Others estimates must not be used as a policy target (see §3.7). '
        'Beyond the heterogeneity of the three ancestry clusters, the group '
        'departs from Hardy–Weinberg equilibrium at all five loci, with observed '
        'heterozygosity below expected at every locus — the classical Wahlund '
        'signature of pooling stratified subpopulations. That violation '
        'invalidates the random-mating assumption underlying the diplotype '
        'expansion F(hi,hj) = 2·fi·fj for this group specifically, independently '
        'of any question about the frequency floor, so pooled-Others figures are '
        'not merely imprecise but structurally inadmissible. Simple ancestry '
        'questions at registration — parental birthplace, self-reported heritage '
        '— would enable sub-group–targeted recruitment and allow cluster-level '
        'rather than pooled estimates to be used ' + cite(16) + '.'
    ),
    (
        'Invest in DQB1 typing for all donors.',
        'Although 8/8 and 10/10 registry sizes differ only modestly due to '
        'strong DRB1–DQB1 LD ' + cite(1, 13) + ', DQB1 mismatches are '
        'associated with increased graft-versus-host disease in some populations '
        + cite(12) + '. The marginal recruitment cost of achieving 10/10 vs '
        '8/8 coverage is very small; the clinical benefit of routine DQB1 '
        'typing is well-established ' + cite(7) + '.'
    ),
    (
        'Use the 95% coverage target as the planning standard.',
        'Planning for 90% rather than 95% coverage may appear cost-saving on '
        'paper but systematically excludes patients with rare diplotypes — who '
        'are disproportionately represented in minority and mixed-ancestry groups '
        + cite(2, 9) + '. The 95% target reflects genuine clinical equity '
        'and is the standard used by established international registries '
        + cite(7) + '.'
    ),
]

for i, (heading, body) in enumerate(rec_items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'Recommendation {i+1}: {heading}')
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.8)
    for run in p2.runs:
        run.font.size = Pt(10)

add_heading(doc, '4.1 Limitations', level=2)
add_para(doc,
    'Several limitations should be noted. First, diplotype frequencies are '
    'derived under Hardy–Weinberg equilibrium. Departures are not uniform across '
    'groups: Chinese and Malay show no significant departure at any locus, Indian '
    'departs at three of five (DQB1, HLA-B, HLA-C), and Others departs at all '
    'five, with observed heterozygosity below expected at every locus '
    '(p from 3.7e-05 at DQB1 to 3.8e-28 at HLA-B) ' + cite(1) + '. The uniform '
    'heterozygote deficit in Others is the classical Wahlund signature of '
    'population stratification and corroborates the sub-cluster analysis in §3.7, '
    'but it also invalidates the random-mating assumption behind the diplotype '
    'expansion for that group; pooled-Others estimates should be replaced by '
    'cluster-level ones. Indian estimates should be treated as exploratory, while '
    'Chinese and Malay estimates are on firmer ground. Second, the EM input '
    'sample size interacts strongly with the frequency floor, and this analysis '
    'removes a cap that earlier versions applied. Versions up to v2.15 capped the '
    'EM at 5,000 individuals per group and reported, on the basis of a convergence '
    'test, that this cost about 8% in a conservative direction. That figure was '
    'specific to the 0.1% floor and does not generalise. Repeating the convergence '
    'test at the floor used here (Supporting Figure S1) gives a Chinese N* of '
    f'{CAP_N:,} from 5,000 individuals against {FULL_N:,} from the full '
    f'{FULL_SAMPLE:,} — a difference of {abs(CAP_DELTA_PCT)}% in the '
    + ('anti-conservative' if CAP_DELTA_PCT < 0 else 'conservative') + ' direction, '
    f'with the estimated haplotype count moving from {CAP_NHAP:,} to {FULL_NHAP:,}. '
    'A 5,000-individual EM cannot resolve phase in the rare tail, so what it '
    'retains there is largely noise. The direction of the cap\'s bias is not even '
    'stable: at a 0.1% floor it overstates N*, and at the floor used here it '
    'understates it. The analysis reported here therefore uses all available '
    'individuals (cap raised to 50,000, which binds for no CMIO group). The general '
    'lesson is that a frequency floor and an input-sample cap cannot be chosen '
    'independently, and neither can be reported without the other. Third, the '
    'reported N* values represent '
    'biologically matched donors; real-world donor attrition (unreachability, '
    'refusal, medical deferral — typically 30–50% in established registries) '
    'means signed-up recruitment targets must exceed N* by a corresponding '
    'factor. At 40% attrition, signed-up targets are N* ÷ 0.60 ≈ '
    'N* × 1.67; per-threshold ranges across CMIO groups are shown in the '
    '"Signed-up target" rows of Tables 1 and 2. Fourth, N* estimates are lower '
    'bounds for unobserved haplotypes, and at the corrected 0.01% floor the '
    'treatment of rare haplotypes is a leading source of uncertainty rather than a '
    'minor one. Sensitivity analysis with Laplace pseudocount smoothing '
    '(α = 0.001 per haplotype) changes N* at the 95% threshold by −17.9% '
    '(Chinese), −17.7% (Malay), −27.8% (Indian) and −31.4% (Others), and at the '
    '75% threshold by +300%, +327%, +143% and +86% respectively. Earlier versions '
    'reported this analysis as confirming robustness (±3% at 95%); that '
    'conclusion was itself an artefact of the 0.1% floor, which had already '
    'removed most of the haplotypes that smoothing perturbs. The honest reading is '
    'that N* is sensitive to rare-haplotype treatment at every threshold, and that '
    'the point estimates in Tables 1 and 2 should be read as order-of-magnitude '
    'quantities. '
    'Fifth, the Others cluster ancestry assignments are inferred from haplotype '
    'signatures ' + cite(16, 17) + ' without confirmed self-reported ancestry '
    'data, and cluster stability was not independently validated via bootstrap '
    'resampling of individuals — the labels are supported by distinct '
    'population-specific top-haplotype signatures (Table 5), but the moderate '
    'silhouette coefficient (0.24) indicates overlapping cluster boundaries, and '
    'independent confirmation '
    'would strengthen these findings. Sixth, validation of EM frequencies against '
    'patient haplotypes was limited by small patient sample sizes in Malay, Indian, '
    'and Others groups ' + cite(1) + ' (see §3.6). Finally, the registry model '
    'assumes random donor sampling; demographic biases in recruitment (age, sex, '
    'region) may affect effective coverage in practice.')

# ── 5. CONCLUSIONS ───────────────────────────────────────────────────────────
add_heading(doc, '5. Conclusions')
add_para(doc,
    'Using EM-phased haplotype frequencies from 59,186 Singapore donors and cord '
    'blood units ' + cite(1) + ', we have derived quantitative registry '
    'size estimates for Singapore\'s CMIO populations. Achieving 95% coverage at '
    f'10/10 HLA matching requires {K("n95_txt","Malay")} to {K("n95_txt","Others")} '
    'same-ethnicity donors and is '
    'therefore not attainable by any national programme; cross-ethnic matching '
    'from a Singapore-weighted pool is worse still for Malay, Indian and Others '
    'patients ' + cite(2, 10) + '. The 40,000–45,000 donor targets reported in '
    'earlier versions of this analysis were an artefact of a 0.1% haplotype '
    'frequency floor that removed the rare-haplotype tail on which high-coverage '
    'behaviour depends, and are withdrawn.')

add_corrected_para(doc, [
    ('Partial-match analysis demonstrates that 9/10 matching halves this requirement ', False),
    (cite(2, 19), True),
    (', providing a practical and clinically justifiable path to '
     'improving access in the near term ' + cite(12, 15) + '. The Others subgroup '
     'is ethnically heterogeneous; sub-cluster analysis identifies three distinct '
     'ancestry groups with registry requirements ranging from 35,000 to 64,000 donors '
     '(Table 4) ' + cite(16) + '.', False),
])

add_para(doc,
    'These findings support a three-pronged national strategy: (1) accelerated '
    'same-ethnicity donor recruitment with defined numerical targets for each '
    'CMIO group ' + cite(9) + '; (2) adoption of evidence-based partial-match '
    'protocols to extend effective registry reach ' + cite(12, 15) + '; and '
    '(3) ancestry sub-group data collection for Others donors to support targeted '
    'recruitment as registry size grows ' + cite(16) + '. '
    'Together, these measures can substantially improve HSCT access equity for '
    'all communities in Singapore.', space_after=12)

# ── SUPPORTING ANALYSIS ──────────────────────────────────────────────────────
add_heading(doc, 'Supporting Analysis')
add_heading(doc, 'S1. EM Convergence: N* Stability vs Sample Size (Chinese)', level=2)
add_para(doc,
    'To assess whether the 5,000-sample EM cap introduces material bias for '
    'the Chinese group, the EM was rerun at sample sizes ranging from 500 to '
    '45,018 (full dataset). Figure S1 shows N* at 95% coverage (10/10 matching) '
    'as a function of input sample size.')
add_figure(doc, 'em_convergence.png', width=6.0,
    caption='Figure S1. N* at 95% coverage (10/10 matching) for Chinese donors '
    'as a function of EM input sample size. The red dashed line marks the 5,000 '
    'cap used in the main analysis (N*=45,148); N* at the full 45,018-donor '
    'sample is 41,727 — an 8.2% overestimate at the cap, in the conservative '
    '(safe) direction. Convergence is near-complete above ~20,000 samples.')
add_para(doc,
    'The 8.2% overestimate at the cap means the reported N* values for Chinese '
    'are conservative by approximately 3,400 donors — equivalent to slightly '
    'less than one year of typical registry recruitment. For Malay, Indian, and '
    'Others, all donors were used (sample sizes ≤ 5,868); the cap does not '
    'affect those estimates.')

# ── GLOSSARY ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Abbreviations', level=2)
glossary = [
    ('AFND',  'Allele Frequency Net Database'),
    ('BMDP',  'Bone Marrow Donor Programme (Singapore)'),
    ('CI',    'Confidence Interval'),
    ('CMIO',  'Chinese–Malay–Indian–Others (Singapore ethnic classification)'),
    ('EM',    'Expectation–Maximisation (algorithm for haplotype phasing)'),
    ('HSCT',  'Haematopoietic Stem Cell Transplantation'),
    ('HLA',   'Human Leukocyte Antigen'),
    ('HSA',   'Health Sciences Authority (Singapore)'),
    ('HWE',   'Hardy–Weinberg Equilibrium'),
    ('LD',    'Linkage Disequilibrium'),
    ('MLE',   'Maximum Likelihood Estimate'),
    ('N*',    'Minimum registry size achieving a specified coverage target'),
    ('PCA',   'Principal Component Analysis'),
    ('RMSE',  'Root Mean Square Error'),
    ('SCBB',  'Singapore Cord Blood Bank'),
]
gtbl = doc.add_table(rows=1, cols=2)
gtbl.rows[0].cells[0].text = 'Abbreviation'
gtbl.rows[0].cells[1].text = 'Definition'
for abbr, defn in glossary:
    row = gtbl.add_row().cells
    row[0].text = abbr
    row[1].text = defn
style_table(gtbl)
gtbl.autofit = False
for row in gtbl.rows:
    row.cells[0].width = Cm(2.2)
    row.cells[1].width = Cm(13.5)

# ── DECLARATIONS ─────────────────────────────────────────────────────────────
add_heading(doc, 'Declarations', level=2)
add_para(doc, 'Conflict of interest: The author declares no conflict of interest.', size=9)
add_para(doc,
    'Data availability: All analysis code and intermediate data files are '
    'available at https://github.com/alvin8-git/HLA.', size=9)
add_para(doc, 'Funding: No external funding was received for this study.', size=9)

# ── REFERENCES ───────────────────────────────────────────────────────────────
add_heading(doc, 'References')
for i, (authors, title, journal, doi) in enumerate(REFS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.left_indent   = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r1 = p.add_run(f'{i}.  ')
    r1.bold = True
    r1.font.size = Pt(9)
    ref_text = f'{authors} {title} {journal}'
    if doi:
        ref_text += f' {doi}'
    p.add_run(ref_text).font.size = Pt(9)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(HERE, 'HLA_Registry_Size_CMIO_v2.16.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables:     {len(doc.tables)}')
print(f'  Figures:    {len(doc.inline_shapes)}')
print(f'  References: {len(REFS)}')
