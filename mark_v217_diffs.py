#!/usr/bin/env python
"""Colour v2.17 so that RED == "differs from v2.15", and nothing else.

The manuscript had accumulated hand-set red runs from an older
"tracked corrections" convention (add_corrected_para), which by v2.17 no longer
corresponded to anything a reader could name. This recomputes the colouring
mechanically: every run is reset to black, then a paragraph- and word-level diff
against the v2.15 .docx paints the genuine differences red.

Run AFTER build_report_v217.py. Idempotent — it always rebuilds colour from the
two documents, so re-running cannot accumulate drift.

    python build_report_v217.py && python mark_v217_diffs.py

Table header rows are left white-on-navy: black or red on a dark fill is
unreadable, and the header change is already visible in the body cells beneath.
"""
import copy
import difflib
import re
import sys

from docx import Document
from docx.shared import RGBColor

# Defaults preserve the original v2.17 usage; pass two paths to override:
#   python mark_v217_diffs.py OLD.docx NEW.docx
OLD = 'HLA_Registry_Size_CMIO_v2.15.docx'
NEW = 'HLA_Registry_Size_CMIO_v2.17.docx'
_args = [a for a in sys.argv[1:] if a != '--selftest']
if len(_args) == 2:
    OLD, NEW = _args

RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

WORD = re.compile(r'\S+|\s+')
DRAWING = (
    '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor',
    '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline',
)


def has_image(par):
    return any(par._p.iter(tag) for tag in DRAWING) and any(
        True for tag in DRAWING for _ in par._p.iter(tag))


# Two v2.15 references were withdrawn by the 2026-08-25 citation audit:
#   [8]  Lim et al.      — no PubMed record exists for it
#   [14] Anasetti et al. — a graft-source trial, cited only for a
#                          matching-level claim it cannot support
# Everything after each shifts down. Renumbering is bookkeeping, not a change
# to the text, so the OLD numbering is mapped forward before diffing and red
# keeps meaning "the claim changed". Empty this out once a rebuilt v2.15
# baseline carries the new numbering.
CITE_DROPPED = {8, 14}
CITE_RENUMBER = {k: k - sum(1 for d in CITE_DROPPED if d < k)
                 for k in range(1, 21) if k not in CITE_DROPPED}


def _remap(nums):
    out = [CITE_RENUMBER.get(k, k) for k in nums if k not in CITE_DROPPED]
    return out or nums          # never let a citation vanish entirely


def normalize_refs(t):
    """Restyle old-text cross-references before diffing, so pure notation or
    numbering changes are not painted as additions.

    Two normalisations: §2.4 -> Section 2.4, and the reference renumbering
    that followed withdrawing old [8].
    """
    t = re.sub(r'§(\d[\d.]*) and §(\d[\d.]*)', r'Sections \1 and \2', t)
    t = re.sub(r'§(?=\d)', 'Section ', t)
    if not CITE_RENUMBER:
        return t

    def cite_sub(m):
        nums = [int(x) for x in re.findall(r'\d+', m.group(1))]
        return '[' + ','.join(str(k) for k in _remap(nums)) + ']'

    t = re.sub(r'\[(\d+(?:\s*,\s*\d+)*)\]', cite_sub, t)

    # reference-list lines ("9.  Aljurf M, ...") carry the number too
    def entry_sub(m):
        k = int(m.group(1))
        return f'{CITE_RENUMBER.get(k, k)}.{m.group(2)}'

    return re.sub(r'^(\d{1,2})\.(\s+[A-Z])', entry_sub, t)


def word_mask(old, new):
    """Per-character bool over `new`: True where it differs from `old`.

    `old` is the RAW v2.15 text. normalize_refs is used for paragraph
    ALIGNMENT only (see main) — masking against normalised text would hide a
    changed citation number, which is a visible difference on the page.
    """
    if old is None:
        return [True] * len(new)
    ot, nt = WORD.findall(old), WORD.findall(new)
    mask = []
    for tag, _i1, _i2, j1, j2 in difflib.SequenceMatcher(
            None, ot, nt, autojunk=False).get_opcodes():
        red = tag != 'equal'
        for tok in nt[j1:j2]:
            mask.extend([red] * len(tok))
    # guard: tokenisation must reproduce the string exactly
    assert len(mask) == len(new), (len(mask), len(new))
    return mask


def recolor(par, mask):
    """Repaint a paragraph, splitting runs at colour boundaries.

    Run-level font properties are carried over by deep-copying each run's
    <w:rPr>, so bold/italic/size/style survive the split; only the colour is
    overwritten.
    """
    runs = par.runs
    if not runs:
        return
    # Skip any paragraph containing a drawing (inline or anchored). iter()
    # searches descendants — findall() only looked at direct children of w:r
    # and missed w:r/w:drawing/wp:inline, so figures were stripped as text.
    if any(next(r._r.iter(tag), None) is not None
           for r in runs for tag in DRAWING):
        return  # picture run — leave alone

    snap, pos = [], 0
    for r in runs:
        snap.append((r.text, r._r.find(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')))
    for r in list(runs):
        r._r.getparent().remove(r._r)

    for text, rpr in snap:
        if not text:
            continue
        seg_start = 0
        while seg_start < len(text):
            want = mask[pos + seg_start] if pos + seg_start < len(mask) else False
            seg_end = seg_start + 1
            while seg_end < len(text):
                nxt = mask[pos + seg_end] if pos + seg_end < len(mask) else False
                if nxt != want:
                    break
                seg_end += 1
            run = par.add_run(text[seg_start:seg_end])
            if rpr is not None:
                old_rpr = run._r.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/'
                    '2006/main}rPr')
                if old_rpr is not None:
                    run._r.remove(old_rpr)
                run._r.insert(0, copy.deepcopy(rpr))
            run.font.color.rgb = RED if want else BLACK
            seg_start = seg_end
        pos += len(text)


def cell_text(c):
    return '\n'.join(p.text for p in c.paragraphs)


def table_sig(t):
    return '\n'.join(' | '.join(c.text.strip() for c in r.cells) for r in t.rows)


def pair_tables(old_tables, new_tables, floor=0.35):
    """Greedy one-to-one match on full-table text similarity.

    Header rows are not distinctive here — the 10/10 and 8/8 tables carry
    identical headers, and matching on them alone pairs the 8/8 table against
    the 10/10 one and floods it with false differences. Whole-table text
    separates them. Anything below `floor` is treated as a table new in v2.17.
    """
    old_sigs = [normalize_refs(table_sig(t)) for t in old_tables]
    new_sigs = [table_sig(t) for t in new_tables]
    scored = sorted(
        ((difflib.SequenceMatcher(None, o, n).ratio(), i, j)
         for i, o in enumerate(old_sigs) for j, n in enumerate(new_sigs)),
        reverse=True)
    pairs, used_o, used_n = {}, set(), set()
    for score, i, j in scored:
        if score < floor or i in used_o or j in used_n:
            continue
        pairs[j] = old_tables[i]
        used_o.add(i)
        used_n.add(j)
    return pairs


def main():
    old_doc, new_doc = Document(OLD), Document(NEW)

    # Normalised text is used ONLY to align paragraphs: it lets the differ see
    # a renumbered reference line as the same line and the withdrawn entries as
    # clean deletes, instead of knocking the whole list out of step. The mask is
    # then computed against the RAW original, so a citation whose number changed
    # still shows red — red means "differs from v2.15 on the page".
    old_norm = [normalize_refs(p.text) for p in old_doc.paragraphs]
    old_raw = [p.text for p in old_doc.paragraphs]
    new_paras = [p.text for p in new_doc.paragraphs]

    masks = [None] * len(new_paras)
    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(
            None, old_norm, new_paras, autojunk=False).get_opcodes():
        if tag == 'equal':
            # aligned, but may still differ before normalisation
            for k, j in enumerate(range(j1, j2)):
                masks[j] = word_mask(old_raw[i1 + k], new_paras[j])
        elif tag == 'insert':
            for j in range(j1, j2):
                masks[j] = [True] * len(new_paras[j])
        elif tag == 'replace':
            olds = old_raw[i1:i2]
            for k, j in enumerate(range(j1, j2)):
                masks[j] = word_mask(olds[k] if k < len(olds) else None,
                                     new_paras[j])

    n_red = 0
    for par, mask in zip(new_doc.paragraphs, masks):
        if mask is None:
            mask = [False] * len(par.text)
        if any(mask):
            n_red += 1
        recolor(par, mask)

    # ── tables: match on header row, then diff body cells ────────────────────
    pairs = pair_tables(old_doc.tables, new_doc.tables)
    t_red = 0
    for tj, t in enumerate(new_doc.tables):
        match = pairs.get(tj)
        for ri, row in enumerate(t.rows):
            if ri == 0:
                continue  # header: white on navy, leave it
            for ci, cell in enumerate(row.cells):
                prev = None
                if match is not None and ri < len(match.rows) \
                        and ci < len(match.rows[ri].cells):
                    # raw, not normalised: a renumbered citation in a cell is
                    # still a visible difference and must show red
                    prev = cell_text(match.rows[ri].cells[ci])
                cur = cell_text(cell)
                mask = word_mask(prev, cur)
                if any(mask):
                    t_red += 1
                off = 0
                for p in cell.paragraphs:
                    recolor(p, mask[off:off + len(p.text)])
                    off += len(p.text) + 1  # + the '\n' join

    new_doc.save(NEW)
    print(f'{NEW}: {n_red}/{len(new_paras)} paragraphs and {t_red} table cells '
          f'carry red (differ from v2.15); everything else black.')


if __name__ == '__main__':
    if '--selftest' in sys.argv:
        assert word_mask('a b c', 'a b c') == [False] * 5
        assert word_mask('a b c', 'a X c') == [False, False, True, False, False]
        assert word_mask(None, 'zz') == [True, True]
        # an inserted word reddens only itself, not its neighbours
        m = word_mask('a c', 'a b c')
        assert m[0] is False and m[2] is True
        # §->Section restyling is not an addition
        nm = lambda a, b: word_mask(normalize_refs(a), b)
        assert not any(nm('see §2.4.', 'see Section 2.4.'))
        assert not any(nm('see §3.6 and §4.1.', 'see Sections 3.6 and 4.1.'))
        # nor is the reference renumbering after old [8] was withdrawn
        assert not any(nm('need [8,9].', 'need [8].'))
        assert not any(nm('diversity [1,8].', 'diversity [1].'))
        assert not any(nm('registries [9].', 'registries [8].'))
        assert not any(nm('mismatch [15].', 'mismatch [13].'))   # after both drops
        assert not any(nm('15.  Pidala J, Lee SJ.', '13.  Pidala J, Lee SJ.'))
        assert not any(nm('LD [1,13].', 'LD [1,12].'))
        assert not any(nm('9.  Aljurf M, Weisdorf D.', '8.  Aljurf M, Weisdorf D.'))
        # a genuinely different claim still reddens
        assert any(nm('halves the size [9].', 'doubles the size [8].'))
        print('mark_v217_diffs: self-check OK')
        sys.exit()
    main()
